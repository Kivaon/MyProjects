import os, sys, time, glob, requests, shutil, re, argparse
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from datetime import datetime
import json
import re
import sys
import os
import time
import math
import tiktoken
import csv
import subprocess
from docx import Document   
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tbox_utils import tbox_log

# Версия программы
VERSION = "3.2.0"
DATE = "2026-05-01"

# Попытка импорта OpenAI
try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

# Попытка импорта tiktoken отдельно
try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
except ImportError:
    TIKTOKEN_AVAILABLE = False

# --- MANIFEST ---
# Previous: v6.10.smart-chunking (2026-04-18) - Smart chunking with Gemini
VERSION = "v6.11.gpt-support"
DATE    = "2026-04-27"
NAME    = os.path.basename(__file__)
META    = {"name": NAME, "version": VERSION}

# Глобальные переменные для подсчета токенов
GPT_TOKEN_STATS = {
    "input_tokens": 0,
    "output_tokens": 0,
    "total_tokens": 0,
    "total_cost": 0.0,
    "chunks_processed": 0,
    "model": ""
}

# Флаг для завершения программы после формирования списка
PROGRAM_SHOULD_EXIT = False

# =============================================================================
# GPT МОДЕЛИ И ПОИСК РАБОЧЕЙ МОДЕЛИ
# =============================================================================

def get_gpt_chunk_limits(gpt_model):
    """
    Определение лимитов токенов для чанков в зависимости от GPT модели
    
    Алгоритм работы:
    1. Определяет максимальное количество токенов для модели
    2. Вычисляет безопасный лимит для входных токенов (70% от максимума)
    3. Возвращает рекомендуемые лимиты для чанков
    
    Args:
        gpt_model (str): Название GPT модели
        
    Returns:
        dict: Словарь с лимитами {
            'max_tokens': максимальные токены модели,
            'input_limit': безопасный лимит входных токенов,
            'output_limit': лимит выходных токенов,
            'recommended_chunk_size': рекомендуемый размер чанка в символах
        }
    """
    # Определяем максимальное количество токенов для каждой модели
    if "3.5" in gpt_model.lower():
        max_tokens = 4096      # gpt-3.5-turbo лимит
    elif "4o" in gpt_model.lower():
        max_tokens = 128000    # gpt-4o лимит (128K context)
    elif "4" in gpt_model.lower():
        max_tokens = 8192      # gpt-4 лимит
    else:
        max_tokens = 4096      # по умолчанию для 3.5
    
    # Вычисляем безопасные лимиты (70% для входа, 30% для выхода)
    input_limit = int(max_tokens * 0.7)
    output_limit = int(max_tokens * 0.3)
    
    # Рекомендуемый размер чанка в символах (примерно 4 символа на токен для английского)
    # Для русского текста коэффициент может быть другим, но используем консервативную оценку
    recommended_chunk_size = input_limit * 3
    
    return {
        'max_tokens': max_tokens,
        'input_limit': input_limit,
        'output_limit': output_limit,
        'recommended_chunk_size': recommended_chunk_size
    }

def get_gemini_chunk_limits(gemini_model):
    """
    Определение лимитов токенов для чанков в зависимости от Gemini модели
    
    Алгоритм работы:
    1. Определяет максимальное количество токенов для модели Gemini
    2. Вычисляет безопасный лимит для входных токенов (70% от максимума)
    3. Возвращает рекомендуемые лимиты для чанков
    
    Args:
        gemini_model (str): Название Gemini модели
        
    Returns:
        dict: Словарь с лимитами {
            'max_tokens': максимальные токены модели,
            'input_limit': безопасный лимит входных токенов,
            'output_limit': лимит выходных токенов,
            'recommended_chunk_size': рекомендуемый размер чанка в символах
        }
    """
    # Определяем максимальное количество токенов для каждой модели Gemini
    if "1.5" in gemini_model.lower() or "gemini-1.5" in gemini_model.lower():
        max_tokens = 32768     # Gemini 1.5 лимит (32K context)
    elif "1.0" in gemini_model.lower() or "gemini-1.0" in gemini_model.lower():
        max_tokens = 32768     # Gemini 1.0 лимит (32K context)
    elif "pro" in gemini_model.lower():
        max_tokens = 2097152   # Gemini 1.5 Pro лимит (2M context)
    else:
        max_tokens = 32768     # по умолчанию для Gemini 1.5
    
    # Вычисляем безопасные лимиты (70% для входа, 30% для выхода)
    input_limit = int(max_tokens * 0.7)
    output_limit = int(max_tokens * 0.3)
    
    # Рекомендуемый размер чанка в символах (примерно 4 символа на токен)
    # Для русского текста используем консервативную оценку
    recommended_chunk_size = min(input_limit * 3, 15000)  # Ограничиваем 15K символов для стабильности
    
    return {
        'max_tokens': max_tokens,
        'input_limit': input_limit,
        'output_limit': output_limit,
        'recommended_chunk_size': recommended_chunk_size
    }

# Список доступных GPT моделей для проверки (в порядке приоритета)
# Модели расположены от самых новых к самым# Статический список GPT-5 моделей для fallback
GPT_MODELS = [
    "gpt-5.5-pro-2026-04-23",   # Самая новая Pro версия
    "gpt-5.5-pro",              # Pro версия
    "gpt-5.5-2026-04-23",       # Конкретная версия
    "gpt-5.5",                  # Базовая GPT-5.5
    "gpt-5-chat-latest",        # Последняя chat версия
    "gpt-5-2025-08-07",         # Конкретная версия
    "gpt-5",                    # Базовая GPT-5
    "gpt-5-mini-2025-08-07",    # Mini версия
    "gpt-5-mini",               # Базовая mini
    "gpt-5-nano-2025-08-07",    # Nano версия
]

def get_available_gpt_models(api_key, conf):
    """
    Получение списка доступных GPT моделей напрямую от OpenAI API
    
    Алгоритм работы:
    1. Запрашивает список всех моделей у OpenAI API
    2. Фильтрует только GPT модели (gpt-*)
    3. Сортирует по приоритету (новые модели первыми)
    4. Исключает экспериментальные и deprecated модели
    
    Args:
        api_key (str): API ключ для OpenAI
        conf (dict): Конфигурация для логирования
        
    Returns:
        list: Список доступных GPT моделей в порядке приоритета
    """
    try:
        tbox_log("🔍 Запрос списка моделей у OpenAI API...", META, "INFO", conf)
        
        client = openai.OpenAI(api_key=api_key)
        models = client.models.list()
        
        # Фильтруем и сортируем модели
        gpt_models = []
        
        # Приоритеты моделей (ТОЛЬКО новые GPT-5)
        priority_keywords = [
            "gpt-5.5",         # GPT-5.5 серии (самые новые)
            "gpt-5",           # GPT-5 серии
            "gpt-5-chat",      # GPT-5 chat версии
        ]
        
        for model in models.data:
            model_id = model.id
            
            # Пропускаем не-GPT модели
            if not model_id.startswith("gpt-"):
                continue
            
            # Пропускаем все НЕ-GPT-5 модели
            if not model_id.startswith("gpt-5"):
                continue
                
            # Пропускаем экспериментальные и deprecated модели
            if any(k in model_id.lower() for k in ["experimental", "deprecated", "preview"]):
                continue
            
            # Добавляем только GPT-5 модели
            gpt_models.append(model_id)
        
        # Сортируем по приоритету
        def get_priority(model_name):
            for i, keyword in enumerate(priority_keywords):
                if keyword in model_name:
                    return i
            return len(priority_keywords)  # В конец списка
        
        gpt_models.sort(key=get_priority)
        
        tbox_log(f"✅ Найдено {len(gpt_models)} GPT моделей", META, "INFO", conf)
        for model in gpt_models[:5]:  # Показываем первые 5
            tbox_log(f"   📋 {model}", META, "INFO", conf)
        
        return gpt_models
        
    except Exception as e:
        tbox_log(f"❌ Ошибка получения списка моделей: {str(e)}", META, "ERROR", conf)
        # Возвращаем стандартный список если API недоступен
        return GPT_MODELS

def detect_model_api_type(api_key, model_id):
    """
    Определяет тип API для модели через тестирование
    
    Args:
        api_key (str): API ключ для OpenAI
        model_id (str): ID модели для тестирования
        
    Returns:
        dict: {'old_api': bool, 'new_api': bool, 'recommended': str}
    """
    client = openai.OpenAI(api_key=api_key)
    
    # Тестируем СТАРЫЙ API
    old_api_works = False
    try:
        response = client.chat.completions.create(
            model=model_id,
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": "Say 'Hello'"}
            ],
            max_tokens=10,
            temperature=0.7
        )
        old_api_works = True
    except:
        pass
    
    # Тестируем НОВЫЙ API
    new_api_works = False
    try:
        response = client.responses.create(
            model=model_id,
            input="Say 'Hello'"
        )
        new_api_works = True
    except:
        pass
    
    # Определяем рекомендуемый API
    if old_api_works and new_api_works:
        recommended = "both"
    elif new_api_works:
        recommended = "new"
    elif old_api_works:
        recommended = "old"
    else:
        recommended = "none"
    
    return {
        'old_api': old_api_works,
        'new_api': new_api_works,
        'recommended': recommended
    }

def find_working_gpt_model(api_key, conf):
    """
    Поиск рабочей GPT модели с динамическим определением API типа
    
    Алгоритм работы:
    1. Получает список доступных моделей от OpenAI API
    2. Для каждой модели определяет тип API через тестирование
    3. Делает тестовый запрос с правильным API
    4. При успешном ответе возвращает название модели и тип API
    5. При ошибке 404 (модель не найдена) продолжает поиск
    6. При ошибке 401 (неверный API ключ) прекращает поиск
    
    Args:
        api_key (str): API ключ для OpenAI
        conf (dict): Конфигурация для логирования
        
    Returns:
        str: Название рабочей модели или None если не найдена
    """
    tbox_log("🔍 Поиск рабочей GPT модели...", META, "INFO", conf)
    
    # Получаем актуальный список моделей
    available_models = get_available_gpt_models(api_key, conf)
    
    # Перебираем модели в порядке приоритета
    for model in available_models:
        try:
            tbox_log(f"🧪 Пробую модель: {model}", META, "INFO", conf)
            
            # Создаем клиент OpenAI с API ключом
            client = openai.OpenAI(api_key=api_key)
            
            # ТОЛЬКО GPT-5 модели используем новый API
            response = client.responses.create(
                model=model,
                input="Say 'Hello'"
            )
            # В GPT-5 моделях ответ через output_text
            test_response = response.output_text
            
            # Если запрос успешен, модель работает
            tbox_log(f"✅ Модель {model} работает!", META, "INFO", conf)
            return model
            
        except openai.NotFoundError as e:
            # Модель не найдена (404) - продолжаем поиск следующей
            tbox_log(f"❌ Модель {model} не найдена", META, "WARN", conf)
            continue
        except openai.AuthenticationError as e:
            # Ошибка аутентификации (401) - API ключ неверный, прекращаем поиск
            tbox_log(f"❌ Ошибка аутентификации для {model}", META, "WARN", conf)
            continue
        except Exception as e:
            tbox_log(f"⚠️ Ошибка с моделью {model}: {str(e)[:50]}", META, "WARN", conf)
            continue
    
    tbox_log("❌ Не найдено рабочих GPT моделей", META, "ERROR", conf)
    return None

def load_tbox_config():
    script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    config_path = os.path.join(script_dir, "_config", "tconfig.txt")
    conf = {}
    if not os.path.exists(config_path): 
        print(f"❌ Ошибка: Конфиг не найден: {config_path}")
        return None
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.split('#')[0].strip()
                if '=' in line:
                    key, val = line.split('=', 1)
                    conf[key.strip()] = val.strip()
        actual_base = conf.get('BASE_DIR', script_dir)
        for key in conf:
            if '${BASE_DIR}' in conf[key]:
                conf[key] = conf[key].replace('${BASE_DIR}', actual_base)
        return conf
    except Exception as e:
        print(f"❌ Ошибка в load_tbox_config: {e}")
        return None

def load_prompts(conf):
    """Загрузка справочника промптов из файла"""
    prompts_dir = conf.get('PROMPTS_DIR', '07_PROMPTS')
    base_dir = conf.get('BASE_DIR', os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    prompts_path = os.path.join(base_dir, prompts_dir, 'prompts.md')
    prompts = {
        'TORAH': "Ты — редактор лекций Рава {author}. Переведи часть {part_num}/{total_parts} на русский.\nСТРОГО СОХРАНЯЙ РАЗМЕТКУ: '#' для заголовков и '**' для выделений.\nСтиль: возвышенный, точный.\nТЕКСТ:\n{chunk}",
        'FICTION': "Переведи художественный текст часть {part_num}/{total_parts} на русский.\nСОХРАНЯЙ стилистику и тон оригинала.\nРазметка: '#' для заголовков, '**' для выделений.\nТЕКСТ:\n{chunk}",
        'GENERIC': "Переведи текст часть {part_num}/{total_parts} на русский.\nТЕКСТ:\n{chunk}"
    }
    if os.path.exists(prompts_path):
        try:
            with open(prompts_path, 'r', encoding='utf-8') as f:
                content = f.read()
            # Парсинг по заголовкам # CODE
            sections = re.split(r'^# (.+)$', content, flags=re.MULTILINE)
            for i in range(1, len(sections), 2):
                code = sections[i].strip()
                text = sections[i+1].strip()
                prompts[code] = text
            #tbox_log(f"Загружено промптов: {list(prompts.keys())}", META, "INFO", conf)
        except Exception as e:
            tbox_log(f"❌ Ошибка загрузки промптов: {e}", META, "ERROR", conf)
    else:
        tbox_log(f"❌ Файл промптов не найден, используем встроенные: {prompts_path}", META, "WARN", conf)
    return prompts

def check_gpt_config(conf):
    """Проверка наличия GPT параметров в конфиге"""
    gpt_model = conf.get('MODEL_GPT', '').strip()
    gpt_key = conf.get('OPENAI_API_KEY', '').strip()
    
    if not gpt_model:
        tbox_log("❌ MODEL_GPT не найден в конфигурации", META, "ERROR", conf)
        return False
    
    if not gpt_key:
        tbox_log("❌ OPENAI_API_KEY не найден в конфигурации", META, "ERROR", conf)
        return False
    
    if not OPENAI_AVAILABLE:
        tbox_log("❌ OpenAI библиотека не установлена. Установите: pip install openai", META, "ERROR", conf)
        return False
    
    if not TIKTOKEN_AVAILABLE:
        tbox_log("❌ tiktoken библиотека не установлена. Установите: pip install tiktoken", META, "WARN", conf)
        tbox_log("❌ Будет использоваться примерная оценка токенов", META, "WARN", conf)
    
    tbox_log(f"GPT конфигурация проверена: модель={gpt_model}", META, "INFO", conf)
    return True

def get_api_ver(model_name):
    """Определяет API версию по названию модели"""
    v_match = re.search(r'(\d+)', model_name)
    v_major = int(v_match.group(1)) if v_match else 1
    if v_major >= 2 or any(x in model_name.lower() for x in ['exp', 'beta']):
        return "v1beta"
    return "v1"

# Глобальные переменные для динамического списка моделей
gemini_fallback_list = None
current_gemini_index = -1

def find_working_model(api_key, conf, exclude_model=None):
    """Автоматический подбор рабочей модели при сбое основной с динамическим списком"""
    global gemini_fallback_list, current_gemini_index
    
    tbox_log("ПОИСК АЛЬТЕРНАТИВНОЙ МОДЕЛИ...", META, "WARN", conf)
    
    # Если динамический список еще не создан, создаем его
    if gemini_fallback_list is None:
        gemini_fallback_list = create_gemini_fallback_list(api_key, conf, exclude_model)
        current_gemini_index = 0
        
        # Если флаг завершения установлен, выходим после создания списка
        if PROGRAM_SHOULD_EXIT:
            sys.exit(0)
        
        if not gemini_fallback_list:
            tbox_log("❌ Не удалось создать список моделей", META, "ERROR", conf)
            return None
        
        tbox_log(f"✅ Обнаружены альтернативные модели: {gemini_fallback_list[:5]}", META, "INFO", conf)
    
    # Возвращаем текущую модель из списка
    if current_gemini_index < len(gemini_fallback_list):
        model = gemini_fallback_list[current_gemini_index]
        tbox_log(f"👉 Выбрана модель: {model} (индекс {current_gemini_index})", META, "INFO", conf)
        return model
    else:
        tbox_log("❌ Список моделей исчерпан", META, "ERROR", conf)
        return None

def create_gemini_fallback_list(api_key, conf, exclude_model=None):
    """Создание динамического списка моделей с фильтрацией и приоритетами"""
    global PROGRAM_SHOULD_EXIT
    def get_gemini_model_priority(model_name):
        """Новый приоритет: версия + тип (lite менее приоритетен)"""
        # Приоритет для latest моделей - самые новые
        if 'latest' in model_name:
            version = 999.0  # Максимальный приоритет для latest
        elif 'preview' in model_name:
            # Для preview моделей ищем версию в названии
            v_match = re.search(r'gemini-(\d+)(?:\.(\d+))?', model_name)
            if v_match:
                major = float(v_match.group(1))
                minor = float(v_match.group(2)) if v_match.group(2) else 0
                version = major + minor / 10
            else:
                version = 1.0  # fallback для preview без версии
        else:
            # Для стабильных моделей
            v_match = re.search(r'gemini-(\d+)(?:\.(\d+))?', model_name)
            if v_match:
                major = float(v_match.group(1))
                minor = float(v_match.group(2)) if v_match.group(2) else 0
                version = major + minor / 10
            else:
                version = 0.0
        
        # Тип модели: flash > flash-lite
        type_penalty = 0
        if 'flash-lite' in model_name:
            type_penalty = 0.5
        elif 'flash' in model_name:
            type_penalty = 0
        elif 'pro' in model_name:
            type_penalty = 0.2
        
        return version - type_penalty
    
        
        
    # 1. Получение списка и параметров
    all_models = []
    model_info = {}  # Храним полную информацию о моделях
    all_stages_data = []  # Храним данные для всех этапов
    
    for api_v in ["v1beta", "v1"]:
        try:
            url = f"https://generativelanguage.googleapis.com/{api_v}/models?key={api_key}"
            r = requests.get(url, timeout=10)
            if r.status_code == 200:
                models_data = r.json().get('models', [])
                for m in models_data:
                    if 'generateContent' in m.get('supportedGenerationMethods', []):
                        model_name = m['name'].split('/')[-1]
                        
                        # Проверяем, не добавляли ли уже эту модель
                        if model_name not in model_info:
                            all_models.append(model_name)
                            
                            # Сохраняем полную информацию о модели
                            model_info[model_name] = {
                                'name': model_name,
                                'display_name': m.get('displayName', ''),
                                'description': m.get('description', ''),
                                'version': m.get('version', ''),
                                'inputTokenLimit': m.get('inputTokenLimit', 0),
                                'outputTokenLimit': m.get('outputTokenLimit', 0),
                                'maxOutputTokens': m.get('maxOutputTokens', 0),
                                'supportedGenerationMethods': m.get('supportedGenerationMethods', []),
                                'api_version': api_v,
                                'has_generate_content': 'generateContent' in m.get('supportedGenerationMethods', []),
                                'is_text_model': False,  # будет определено позже
                                'input_limit_zero': False,  # будет определено позже
                                'post_status': None,  # будет определено позже
                                'post_response': None,  # будет определено позже
                                'is_working': False  # будет определено позже
                            }
                
                # print(f"    ✅ {api_v} - найдено {len(models_data)} моделей")
            else:
                print(f"    ❌ {api_v} - status {r.status_code}")
        except Exception as e:
            print(f"    ❌ {api_v} - ошибка: {e}")
            continue
    
    # Этап 2: Комплексная фильтрация
    
    # Список исключений (неподходящие модели по имени)
    exception_patterns = [
        'embedding', 'image', 'vision', 'video', 'audio', 'multimodal', 
        'code', 'chat', 'imager', 'imagen', 'aqa', 'retrieval',
        'robotics', 'clip', 'customtools', 'lyria'
    ]
    
    filtered_models = []
    excluded_by_input_limit = []
    excluded_by_output_limit = []
    excluded_by_name = []
    excluded_by_text = []
    
    for model in all_models:
        if model not in model_info:
            continue
            
        info = model_info[model]
        input_limit = info.get('inputTokenLimit', 0)
        output_limit = info.get('outputTokenLimit', 0)
        
        # 1. Фильтр по inputTokenLimit > 200000
        if input_limit <= 200000:
            excluded_by_input_limit.append(model)
            continue
        
        # 2. Фильтр по outputTokenLimit > 16000
        if output_limit <= 16000:
            excluded_by_output_limit.append(model)
            continue
        
        # 3. Фильтр по исключениям в имени
        model_lower = model.lower()
        if any(pattern in model_lower for pattern in exception_patterns):
            excluded_by_name.append(model)
            continue
        
        # 4. Фильтр по текстовым моделям (generateContent + не изображение)
        if not info.get('has_generate_content', False):
            excluded_by_text.append(model)
            continue
            
        # Проверяем, что это не модель для изображений
        description = info.get('description', '').lower()
        display_name = info.get('display_name', '').lower()
        if any(word in description + display_name for word in ['image', 'vision', 'multimodal']):
            excluded_by_text.append(model)
            continue
        
        # Если все проверки пройдены
        filtered_models.append(model)
        info['is_text_model'] = True  # Устанавливаем флаг текстовой модели
    
    # Вывод таблицы после этапа 2
    stage2_data = [model_info[m] for m in filtered_models if m in model_info]
    all_stages_data.append(("ЭТАП 2: КОМПЛЕКСНАЯ ФИЛЬТРАЦИЯ", stage2_data))
    
    # Продолжаем с отфильтрованными моделями       
    # Этап 3: POST запросы для тестирования моделей
    working_models = []
    failed_models = []
    
    for i, model in enumerate(filtered_models):
        # Тестовый POST запрос
        try:
            api_ver = get_api_ver(model)
            url = f"https://generativelanguage.googleapis.com/{api_ver}/models/{model}:generateContent?key={api_key}"
            
            test_prompt = "Hello, please respond with a simple greeting."
            
            r = requests.post(url, json={"contents": [{"parts": [{"text": test_prompt}]}]}, timeout=30)
            
            # Сохраняем результаты теста
            if model in model_info:
                model_info[model]['post_status'] = r.status_code
                model_info[model]['post_response'] = r.text[:100] if r.text else ''
                
                if r.status_code == 200:
                    # Проверяем наличие текстового ответа
                    try:
                        response_data = r.json()
                        if (response_data and 
                            'candidates' in response_data and 
                            len(response_data['candidates']) > 0 and
                            'content' in response_data['candidates'][0] and
                            'parts' in response_data['candidates'][0]['content'] and
                            len(response_data['candidates'][0]['content']['parts']) > 0 and
                            'text' in response_data['candidates'][0]['content']['parts'][0]):
                            
                            content = response_data['candidates'][0]['content']['parts'][0]['text']
                            model_info[model]['is_working'] = True
                            working_models.append(model)
                        else:
                            model_info[model]['is_working'] = False
                            failed_models.append(model)
                    except (KeyError, IndexError, TypeError) as e:
                        model_info[model]['is_working'] = False
                        failed_models.append(model)
                else:
                    model_info[model]['is_working'] = False
                    failed_models.append(model)
            
        except Exception as e:
            if model in model_info:
                model_info[model]['post_status'] = 'ERROR'
                model_info[model]['post_response'] = str(e)[:50]
                model_info[model]['is_working'] = False
            failed_models.append(model)
        
        # Пауза 2 секунды между запросами (кроме последнего)
        if i < len(filtered_models) - 1:
            time.sleep(2)
    
    # Вывод таблицы после этапа 3
    stage3_data = [model_info[m] for m in working_models if m in model_info]
    all_stages_data.append(("ЭТАП 3: POST ЗАПРОСЫ - ТЕСТИРОВАНИЕ", stage3_data))
    
    # Дубликаты уже отфильтрованы при получении моделей из API
    unique_models = working_models.copy()
    
    # Вывод таблицы после этапа 4
    stage4_data = [model_info[m] for m in unique_models if m in model_info]
    all_stages_data.append(("ЭТАП 4: ТЕСТИРОВАНИЕ МОДЕЛЕЙ", stage4_data))
    
    # 6. Исключаем указанную модель если нужно
    if exclude_model and exclude_model in unique_models:
        unique_models.remove(exclude_model)
    
    # 7. Сортировка по приоритету модели, внутри - по типу (lite ниже)
    sorted_models = sorted(unique_models, key=get_gemini_model_priority, reverse=True)
    
    return sorted_models[:10]  # Возвращаем топ-10 моделей
    
    valid_models = []
    zero_input_models = []
    
    # Проверяем первые 5 моделей для отладки
    print("🔍 Отладка - первые 5 моделей:")
    for i, model in enumerate(all_models[:5]):
        if model in model_info:
            info = model_info[model]
            input_limit = info.get('inputTokenLimit', 0)
            output_limit = info.get('outputTokenLimit', 0)
            max_output = info.get('maxOutputTokens', 0)
            print(f"  {i+1}. {model}: input={input_limit}, output={output_limit}, max={max_output}")
        else:
            print(f"⚠️  {i+1}. {model}: НЕ НАЙДЕНА в model_info")
    
    print(f"🔄 Начинаем фильтрацию {len(all_models)} моделей...")
    
    # Добавляем прогресс-бар для отладки
    for i, model in enumerate(all_models):
        if i % 10 == 0:  # Каждые 10 моделей выводим прогресс
            print(f"  🔄 Обработка модели {i+1}/{len(all_models)}: {model}")
        
        if model in model_info:
            input_limit = model_info[model].get('inputTokenLimit', 0)
            model_info[model]['input_limit_zero'] = (input_limit == 0)
            
            if input_limit == 0:
                zero_input_models.append(model)
            else:
                valid_models.append(model)
        else:
            print(f"  ❌ Модель {model} не найдена в model_info!")
    
    print(f"✅ Фильтрация завершена!")
    print(f"🔍 DEBUG: После фильтрации - valid={len(valid_models)}, zero={len(zero_input_models)}")
    
    print(f"❌ Моделей с inputTokenLimit = 0: {len(zero_input_models)}")
    if zero_input_models:
        print(f"    🚫 {zero_input_models[:5]}...")
    
    # Вывод таблицы после этапа 2: Фильтрация по inputTokenLimit
    print("\n" + "="*80)
    print("ЭТАП 2: ФИЛЬТРАЦИЯ ПО INPUT TOKEN LIMIT")
    print("="*80)
    stage2_data = [model_info[m] for m in valid_models if m in model_info]
    print_models_table(stage2_data)
    all_stages_data.append(("ЭТАП 2: ФИЛЬТРАЦИЯ ПО INPUT TOKEN LIMIT", stage2_data))
    
    # Фильтрация текстовых моделей уже сделана в Этапе 2 через exception_patterns
    # Все модели в filtered_models уже являются текстовыми
    
    # 3. Фильтрация рабочих моделей - проверяем реальным запросом
    working_models = []
    for model in filtered_models:
        # Проверяем доступность модели тестовым запросом
        try:
            api_v = get_api_ver(model)
            test_url = f"https://generativelanguage.googleapis.com/{api_v}/models/{model}:generateContent?key={api_key}"
            tr = requests.post(test_url, json={"contents": [{"parts": [{"text": "test"}]}]}, timeout=5)
            
            # Сохраняем результаты теста
            if model in model_info:
                model_info[model]['post_status'] = tr.status_code
                
            if tr.status_code == 200:
                # Детальная проверка - есть ли текстовый контент в ответе
                response_data = tr.json()
                if 'candidates' in response_data and len(response_data['candidates']) > 0:
                    candidate = response_data['candidates'][0]
                    if 'content' in candidate and 'parts' in candidate['content'] and len(candidate['content']['parts']) > 0:
                        if 'text' in candidate['content']['parts'][0]:
                            working_models.append(model)
                            if model in model_info:
                                model_info[model]['is_working'] = True
                                model_info[model]['post_response'] = 'working'
                            print(f"    ✅ {model} - работает")
                        else:
                            if model in model_info:
                                model_info[model]['post_response'] = 'no_text'
                            print(f"    ❌ {model} - нет текста в ответе")
                    else:
                        if model in model_info:
                            model_info[model]['post_response'] = 'no_content_parts'
                        print(f"    ❌ {model} - нет content/parts")
                else:
                    if model in model_info:
                        model_info[model]['post_response'] = 'no_candidates'
                    print(f"    ❌ {model} - нет candidates")
            elif tr.status_code == 429:
                # Превышен лимит запросов
                try:
                    error_data = tr.json()
                    error_msg = error_data.get('error', {}).get('message', 'Лимит превышен')
                    if model in model_info:
                        model_info[model]['post_response'] = f'429: {error_msg}'
                    print(f"    ❌ {model} - лимит запросов: {error_msg}")
                except:
                    if model in model_info:
                        model_info[model]['post_response'] = '429: лимит превышен'
                    print(f"    ❌ {model} - лимит запросов (429)")
            elif tr.status_code == 404:
                if model in model_info:
                    model_info[model]['post_response'] = '404: модель не найдена'
                print(f"    ❌ {model} - модель не найдена (404)")
            else:
                if model in model_info:
                    model_info[model]['post_response'] = f'{tr.status_code}: ошибка'
                print(f"    ❌ {model} - статус {tr.status_code}")
        except Exception as e:
            if model in model_info:
                model_info[model]['post_response'] = f'exception: {str(e)[:50]}'
            continue
    print(f"✅ РАБОЧИЕ МОДЕЛИ (test запрос): {len(working_models)}")
    
    # Вывод таблицы после этапа 4: Тест POST запросов
    print("\n" + "="*80)
    print("ЭТАП 4: ТЕСТ POST ЗАПРОСОВ")
    print("="*80)
    stage4_data = [model_info[m] for m in text_models if m in model_info]
    print_models_table(stage4_data)
    all_stages_data.append(("ЭТАП 4: ТЕСТ POST ЗАПРОСОВ", stage4_data))
    
    # Дубликаты уже отфильтрованы при получении моделей из API
    unique_models = working_models.copy()
    print(f"🔀 Уникальных моделей: {len(unique_models)}")
    
    # Вывод таблицы после этапа 4
    stage4_data = [model_info[m] for m in unique_models if m in model_info]
    print_models_table(stage4_data)
    all_stages_data.append(("ЭТАП 4: ТЕСТИРОВАНИЕ МОДЕЛЕЙ", stage4_data))
    
    # 6. Исключаем указанную модель если нужно
    if exclude_model and exclude_model in unique_models:
        unique_models.remove(exclude_model)
        print(f"🗑️ Исключена {exclude_model}: {len(unique_models)} моделей")
    
    # 7. Сортировка по приоритету модели, внутри - по типу (lite ниже)
    sorted_models = sorted(unique_models, key=get_gemini_model_priority, reverse=True)
    print(f"🔄 ОТСОРТИРОВАННЫЕ: {sorted_models[:5]}...")
    
    # Вывод финальной таблицы
    print("\n" + "="*80)
    print("ЭТАП 6: ФИНАЛЬНЫЙ СПИСОК (ОТСОРТИРОВАННЫЙ)")
    print("="*80)
    stage6_data = [model_info[m] for m in sorted_models if m in model_info]
    print_models_table(stage6_data)
    all_stages_data.append(("ЭТАП 6: ФИНАЛЬНЫЙ СПИСОК (ОТСОРТИРОВАННЫЙ)", stage6_data))
    
    # Сохранение всех этапов в Excel файл
    print(f"\n📁 Сохранение всех этапов в Excel файл...")
    excel_file = save_models_to_excel(all_stages_data)
    print(f"💾 Сохранено: {excel_file}")
    
    # Открытие файла в Excel
    open_excel_file(excel_file)
    
    # Завершение работы после формирования списка
    print(f"\n🎉 Динамический список сформирован!")
    print(f"🎯 Рабочие модели: {sorted_models[:10]}")
    print(f"📂 Результаты сохранены в: {excel_file}")
    
    # Временно убираем завершение, чтобы создать новый CSV с правильными полями
    # PROGRAM_SHOULD_EXIT = True  # Устанавливаем флаг для завершения
    
    return sorted_models[:10]  # Ограничить топ-10 моделей

def print_models_table(models_info, stage_name=""):
    """Вывод таблицы с параметрами моделей"""
    if not models_info:
        print("  📭 Нет моделей для отображения")
        return
    
    # Заголовки таблицы
    headers = [
        "Название", "Display Name", "Описание", "Версия", 
        "Input Limit", "Output Limit", "Max Output", "Methods",
        "API", "GenContent", "Text", "Input=0", "POST Status", "Response", "Working"
    ]
    
    # Форматирование строк
    rows = []
    for model in models_info:
        # Безопасное получение полей для избежания NoneType ошибок
        methods = model.get('supportedGenerationMethods')
        methods_count = len(methods) if methods else 0
        
        row = [
            str(model.get('name', ''))[:20],
            str(model.get('display_name', ''))[:25],
            str(model.get('description', ''))[:30],
            str(model.get('version', '')),
            str(model.get('inputTokenLimit', 0)),
            str(model.get('outputTokenLimit', 0)),
            str(model.get('maxOutputTokens', 0)),
            str(methods_count),
            str(model.get('api_version', '')),
            '✓' if model.get('has_generate_content', False) else '✗',
            '✓' if model.get('is_text_model', False) else '✗',
            '✓' if model.get('input_limit_zero', False) else '✗',
            'n/a' if model.get('post_status') is None else str(model.get('post_status', '')),
            'n/a' if model.get('post_response') is None else str(model.get('post_response', ''))[:15],
            'n/a' if model.get('is_working') is None else ('✓' if model.get('is_working', False) else '✗')
        ]
        rows.append(row)
    
    # Вывод таблицы
    print(f"{'Название':<20} {'Display':<25} {'Описание':<30} {'Версия':<8} {'Input':<8} {'Output':<8} {'MaxOut':<8} {'Meth':<4} {'API':<6} {'Gen':<3} {'Txt':<3} {'In0':<3} {'POST':<6} {'Response':<15} {'Work':<5}")
    print("=" * 170)
    
    for row in rows:
        print(f"{row[0]:<20} {row[1]:<25} {row[2]:<30} {row[3]:<8} {row[4]:<8} {row[5]:<8} {row[6]:<8} {row[7]:<4} {row[8]:<6} {row[9]:<3} {row[10]:<3} {row[11]:<3} {row[12]:<6} {row[13]:<15} {row[14]:<5}")
    
    print(f"\nВсего: {len(models_info)} моделей")

def save_models_to_excel(all_stages_data, filename=None):
    """Сохранение всех этапов в Excel файл"""
    if filename is None:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"gemini_models_all_stages_{timestamp}.csv"
    
    with open(filename, 'w', newline='', encoding='utf-8-sig') as f:
        for stage_name, models_info in all_stages_data:
            # Заголовок этапа
            f.write(f"\n{'='*80}\n")
            f.write(f"# {stage_name}\n")
            f.write(f"{'='*80}\n\n")
            
            if not models_info:
                f.write("Нет данных для этого этапа\n\n")
                continue
            
            # Заголовки CSV
            headers = [
                "Название", "Display Name", "Описание", "Версия", 
                "Input Limit", "Output Limit", "Max Output", "Methods",
                "API", "GenContent", "Text", "Input=0", "POST Status", "Response", "Working"
            ]
            writer = csv.DictWriter(f, fieldnames=headers)
            writer.writeheader()
            
            # Данные моделей
            for model in models_info:
                row = {
                    "Название": model.get('name', ''),
                    "Display Name": model.get('display_name', ''),
                    "Описание": model.get('description', ''),
                    "Версия": model.get('version', ''),
                    "Input Limit": model.get('inputTokenLimit', 0),
                    "Output Limit": model.get('outputTokenLimit', 0),
                    "Max Output": model.get('maxOutputTokens', 0),
                    "Methods": ', '.join(model.get('supportedGenerationMethods', [])),
                    "API": model.get('api_version', ''),
                    "GenContent": '✓' if model.get('has_generate_content', False) else '✗',
                    "Text": '✓' if model.get('is_text_model', False) else '✗',
                    "Input=0": '✓' if model.get('input_limit_zero', False) else '✗',
                    "POST Status": 'n/a' if model.get('post_status') is None else model.get('post_status', ''),
                    "Response": 'n/a' if model.get('post_response') is None else model.get('post_response', ''),
                    "Working": 'n/a' if model.get('is_working') is None else ('✓' if model.get('is_working', False) else '✗')
                }
                writer.writerow(row)
            
            f.write(f"\nВсего моделей: {len(models_info)}\n\n")
    
    return filename

def open_excel_file(filename):
    """Открытие файла в Excel"""
    try:
        if sys.platform == "darwin":  # macOS
            subprocess.run(["open", filename])
        elif sys.platform == "win32":  # Windows
            os.startfile(filename)
        else:  # Linux
            subprocess.run(["xdg-open", filename])
        print(f"📂 Файл открыт в Excel: {filename}")
    except Exception as e:
        print(f"❌ Не удалось открыть файл: {e}")

def count_tokens(text, model="gpt-3.5-turbo"):
    """Подсчет токенов для текста"""
    try:
        encoding = tiktoken.encoding_for_model(model)
        return len(encoding.encode(text))
    except:
        # Если tiktoken недоступен, примерная оценка
        return len(text) // 4

def print_gpt_token_stats(conf):
    """Вывод итоговой статистики по токенам GPT"""
    global GPT_TOKEN_STATS
    
    if GPT_TOKEN_STATS["chunks_processed"] > 0:
        tbox_log("=" * 60, META, "INFO", conf)
        tbox_log("📊 GPT СТАТИСТИКА ПЕРЕВОДА", META, "INFO", conf)
        tbox_log(f"🤖 Модель: {GPT_TOKEN_STATS['model']}", META, "INFO", conf)
        tbox_log(f"📦 Обработано чанков: {GPT_TOKEN_STATS['chunks_processed']}", META, "INFO", conf)
        tbox_log(f"📥 Входные токены: {GPT_TOKEN_STATS['input_tokens']:,}", META, "INFO", conf)
        tbox_log(f"📤 Выходные токены: {GPT_TOKEN_STATS['output_tokens']:,}", META, "INFO", conf)
        tbox_log(f"💰 Всего токенов: {GPT_TOKEN_STATS['total_tokens']:,}", META, "INFO", conf)
        tbox_log(f"💵 Общая стоимость: ${GPT_TOKEN_STATS['total_cost']:.6f}", META, "INFO", conf)
        
        # Дополнительная статистика
        if GPT_TOKEN_STATS["chunks_processed"] > 0:
            avg_input = GPT_TOKEN_STATS["input_tokens"] / GPT_TOKEN_STATS["chunks_processed"]
            avg_output = GPT_TOKEN_STATS["output_tokens"] / GPT_TOKEN_STATS["chunks_processed"]
            avg_cost = GPT_TOKEN_STATS["total_cost"] / GPT_TOKEN_STATS["chunks_processed"]
            
            tbox_log(f"📈 Среднее входных токенов/чанк: {avg_input:.1f}", META, "INFO", conf)
            tbox_log(f"📈 Среднее выходных токенов/чанк: {avg_output:.1f}", META, "INFO", conf)
            tbox_log(f"📈 Средняя стоимость/чанк: ${avg_cost:.6f}", META, "INFO", conf)
        
        tbox_log("=" * 60, META, "INFO", conf)

def translate_with_gpt(chunk, part_num, total_parts, author, conf, prompts, prompt_code, include_original):
    """
    Перевод через OpenAI GPT API с retry-логикой и накоплением токенов
    
    Алгоритм работы:
    1. Готовит промпт для перевода с учетом автора и части
    2. Подсчитывает токены для成本-расчета
    3. Внешний цикл while True позволяет переключиться на другую модель
    4. Внутренний цикл for attempt in range(3) делает 3 попытки для каждой модели
    5. При ошибке 404 (модель не найдена) ищет рабочую модель и переключается
    6. При ошибке 401 (неверный ключ) прекращает попытки
    7. При других ошибках делает retry с задержками
    8. Накапливает статистику токенов и стоимости
    
    Args:
        chunk (str): Текст для перевода
        part_num (int): Номер текущей части
        total_parts (int): Общее количество частей
        author (str): Автор текста
        conf (dict): Конфигурация
        prompts (dict): Справочник промптов
        prompt_code (str): Код промпта
        include_original (bool): Включать ли оригинал в перевод
        
    Returns:
        str: Переведенный текст с метаданными
    """
    global GPT_TOKEN_STATS
    
    # Получаем API ключ из конфигурации
    gpt_key = conf.get('OPENAI_API_KEY', '').strip()
    
    # Находим рабочую модель (если еще не найдена)
    gpt_model = conf.get('MODEL_GPT', '').strip()
    
    # Определяем нужно ли искать модель
    need_search_model = False
    if not gpt_model:
        need_search_model = True
    
    # Если используется --gptlast, ищем последнюю модель
    if conf.get('use_gpt_last') == True:
        need_search_model = True
        tbox_log("🔍 Поиск последней GPT модели (--gptlast)...", META, "INFO", conf)
    elif conf.get('use_gpt_last') == False:
        # Для --gpt не ищем модель, используем из конфига
        need_search_model = False
        tbox_log(f"🎯 Используем модель из конфига: {gpt_model}", META, "INFO", conf)
    
    if need_search_model:
        tbox_log("🔍 Поиск рабочей GPT модели для перевода...", META, "INFO", conf)
        gpt_model = find_working_gpt_model(gpt_key, conf)
        if not gpt_model:
            raise Exception("Не найдено рабочих GPT моделей")
        # Обновляем модель в конфигурации
        conf['MODEL_GPT'] = gpt_model
        tbox_log(f"👉 Выбрана модель: {gpt_model}", META, "INFO", conf)
    
    # Инициализация статистики для первого чанка
    if GPT_TOKEN_STATS["chunks_processed"] == 0:
        GPT_TOKEN_STATS["model"] = gpt_model
    
    # Получаем лимиты для найденной модели
    chunk_limits = get_gpt_chunk_limits(gpt_model)
    
    # Получить базовый промпт из справочника или использовать GENERIC по умолчанию
    base_prompt = prompts.get(prompt_code, prompts.get('GENERIC', 'Переведи на русский: {chunk}'))
    
    # Подставить плейсхолдеры в промпт
    prompt = base_prompt.format(
        author=author,
        part_num=part_num,
        total_parts=total_parts,
        chunk=chunk
    )
    
    # Если требуется, добавляем инструкцию о включении оригинала
    if include_original:
        prompt += "\n\nВКЛЮЧАЙ ОРИГИНАЛЬНЫЙ ТЕКСТ ЦИТАТ В СКОБКАХ ПОСЛЕ ПЕРЕВОДА."
    
    # Подсчет входных токенов для成本-расчета
    system_msg = "Переведи текст на русский язык. Ты профессиональный переводчик. Переводи текст точно, сохраняя форматирование и смысл."
    input_tokens = count_tokens(system_msg + prompt, gpt_model)
    
    # Логирование для отладки проблем с первым чанком
    tbox_log("=" * 60, META, "INFO", CONF)
    tbox_log(f"ℹ️ Чанк {part_num}/{total_parts}: {len(chunk)} символов, осталось {len(text) - sum(len(c) for c in chunks)} символов", META, "INFO", conf)
    tbox_log(f"📊 Токенов: {input_tokens} (вход)", META, "INFO", conf)
    tbox_log(f"🎯 Модель: {gpt_model}", META, "INFO", conf)
    tbox_log(f"ℹ️ Лимиты модели: {chunk_limits['input_limit']} входных / {chunk_limits['output_limit']} выходных токенов", META, "INFO", conf)
    tbox_log(f"📐 Рекомендуемый размер чанка: {chunk_limits['recommended_chunk_size']} символов", META, "INFO", conf)
    
    # Проверяем, не превышает ли чанк безопасные лимиты
    if len(chunk) > chunk_limits['recommended_chunk_size']:
        tbox_log(f"⚠️ Чанк превышает рекомендуемый размер! ({len(chunk)} > {chunk_limits['recommended_chunk_size']})", META, "WARN", conf)
        # Здесь можно добавить логику разбиения чанка на более мелкие части
    elif input_tokens > chunk_limits['input_limit']:
        tbox_log(f"⚠️ Чанк превышает лимит входных токенов! ({input_tokens} > {chunk_limits['input_limit']})", META, "WARN", conf)
    
    # Показываем первые 100 символов для анализа контента
    preview = chunk[:100].replace('\n', ' ').strip()
    tbox_log(f"📖 Предпросмотр: {preview}...", META, "INFO", conf)
    
    # Метаданные для идентификации части перевода
    header_marker = f"\n\n--- [PART {part_num}/{total_parts} | GPT: {gpt_model}] ---\n"
    
    # =============================================================================
    # ОСНОВНОЙ ЦИКЛ ПЕРЕВОДА С ПОДДЕРЖКОЙ ПЕРЕКЛЮЧЕНИЯ МОДЕЛЕЙ
    # =============================================================================
    
    # Внешний цикл while True позволяет перезапустить перевод с новой моделью
    # Это нужно для автоматического переключения при ошибке 404 (модель не найдена)
    while True:
        # Внутренний цикл делает 3 попытки для текущей модели
        for attempt in range(3):
            try:
                tbox_log(f"🔄 Перевод части {part_num}/{total_parts}...Попытка {attempt+1}/3 ({gpt_model})", META, "INFO", conf)
                
                # Создаем клиент OpenAI с API ключом
                client = openai.OpenAI(api_key=gpt_key)
                
                # Используем динамические лимиты для текущей модели
                max_tokens = chunk_limits['output_limit']
                
                # =============================================================================
                # ВЫПОЛНЕНИЕ ЗАПРОСА К GPT API
                # =============================================================================
                
                # ТОЛЬКО НОВЫЙ API для GPT-5 моделей
                response = client.responses.create(
                    model=gpt_model,
                    input=prompt  # Простой текстовый ввод
                )
                # В GPT-5 моделях ответ через output_text
                translation = response.output_text
                output_tokens = count_tokens(translation, gpt_model)
                
                # Логирование ответа для анализа
                tbox_log(f"✅ Ответ GPT получен: {len(translation)} символов, {output_tokens} токенов", META, "INFO", conf)
                
                # Показываем первые 100 символов ответа
                response_preview = translation[:100].replace('\n', ' ').strip()
                tbox_log(f"📖 Предпросмотр ответа: {response_preview}...", META, "INFO", conf)
                
                # =============================================================================
                # РАСЧЕТ СТОИМОСТИ И НАКОПЛЕНИЕ СТАТИСТИКИ
                # =============================================================================
                
                # Расчет стоимости в зависимости от модели (цены за 1 токен)
                if "3.5" in gpt_model.lower():
                    # gpt-3.5-turbo: $0.0000005 входных, $0.0000015 выходных
                    input_cost = input_tokens * 0.0000005
                    output_cost = output_tokens * 0.0000015
                elif "4o-mini" in gpt_model.lower():
                    # gpt-4o-mini: $0.00015 входных, $0.0006 выходных
                    input_cost = input_tokens * 0.00000015
                    output_cost = output_tokens * 0.0000006
                elif "4o" in gpt_model.lower():
                    # gpt-4o: $0.005 входных, $0.015 выходных
                    input_cost = input_tokens * 0.000005
                    output_cost = output_tokens * 0.000015
                elif "4" in gpt_model.lower():
                    # gpt-4: $0.00003 входных, $0.00006 выходных
                    input_cost = input_tokens * 0.00003
                    output_cost = output_tokens * 0.00006
                elif "5" in gpt_model.lower():
                    # gpt-5.5-pro: $0.0000005 входных, $0.0000015 выходных (такие же как 3.5)
                    input_cost = input_tokens * 0.0000005
                    output_cost = output_tokens * 0.0000015
                else:
                    # По умолчанию цены как для 3.5
                    input_cost = input_tokens * 0.0000005
                    output_cost = output_tokens * 0.0000015
                
                total_cost = input_cost + output_cost
                
                # Накопление глобальной статистики для итогового отчета
                GPT_TOKEN_STATS["input_tokens"] += input_tokens
                GPT_TOKEN_STATS["output_tokens"] += output_tokens
                GPT_TOKEN_STATS["total_tokens"] += input_tokens + output_tokens
                GPT_TOKEN_STATS["total_cost"] += total_cost
                GPT_TOKEN_STATS["chunks_processed"] += 1
                
                tbox_log(f"✅ Часть {part_num}/{total_parts} переведена через GPT", META, "INFO", conf)
                return header_marker + translation
            
            except openai.RateLimitError as e:
                # 429 ошибка - превышен лимит запросов (RPM или RPD)
                tbox_log(f"⚠️ GPT Rate Limit на попытке {attempt+1}/3: {str(e)}", META, "WARN", conf)
                if attempt < 2:
                    # Увеличиваем задержку с каждой попыткой: 15, 20, 25 секунд
                    delay = 15 + (5 * attempt)
                    tbox_log(f"⏳ Пауза {delay} сек перед повтором...", META, "INFO", conf)
                    time.sleep(delay)
                    continue
                else:
                    raise Exception(f"GPT лимит запросов превышен после 3 попыток: {str(e)}")
                
            except openai.APIError as e:
                # Общая API ошибка - анализируем статус код
                error_msg = str(e)
                
                if "401" in error_msg or "invalid_api_key" in error_msg:
                    # 401 - неверный API ключ, не retry (критическая ошибка)
                    tbox_log(f"❌ GPT ошибка 401: неверный API ключ", META, "ERROR", conf)
                    raise Exception(f"❗️ GPT ошибка 401: неверный API ключ")
                    
                elif "404" in error_msg or "model_not_found" in error_msg:
                    # 404 - модель не найдена, ищем рабочую модель
                    tbox_log(f"❌ GPT модель не найдена: {gpt_model}", META, "ERROR", conf)
                    new_model = find_working_gpt_model(gpt_key, conf)
                    if new_model and new_model != gpt_model:
                        tbox_log(f"🔄 Найдена рабочая модель: {new_model}", META, "INFO", conf)
                        gpt_model = new_model  # Обновляем локальную модель
                        conf['MODEL_GPT'] = new_model  # Обновляем глобальную конфигурацию
                        tbox_log(f"🌐 Глобальная модель обновлена на: {new_model}", META, "INFO", conf)
                        break  # Выходим из внутреннего цикла, продолжаем внешний с новой моделью
                    else:
                        # Другие API ошибки (500, 503, etc.) - retry с задержками
                        tbox_log(f"⚠️ GPT API ошибка на попытке {attempt+1}/3: {str(e)}", META, "WARN", conf)
                        if attempt < 2:
                            # Увеличиваем задержку: 10, 13, 16 секунд
                            delay = 10 + (3 * attempt)
                            tbox_log(f"⏳ Пауза {delay} сек перед повтором...", META, "INFO", conf)
                            time.sleep(delay)
                            continue
                        else:
                            raise Exception(f"🔴 GPT API ошибка после 3 попыток: {str(e)}")
                    
            except openai.AuthenticationError as e:
                # Ошибка аутентификации - не повторять
                tbox_log(f"❌ GPT ошибка аутентификации: {str(e)}", META, "ERROR", conf)
                raise Exception(f"GPT ошибка аутентификации: {str(e)}")
                
            except openai.NotFoundError as e:
                # Модель не найдена - ищем рабочую модель
                tbox_log(f"❌ GPT модель не найдена: {gpt_model}", META, "ERROR", conf)
                if attempt == 2:  # На последней попытке ищем замену
                    new_model = find_working_gpt_model(gpt_key, conf)
                    if new_model and new_model != gpt_model:
                        tbox_log(f"🔄 Найдена рабочая модель: {new_model}", META, "INFO", conf)
                        gpt_model = new_model  # Обновляем модель
                        break  # Прерываем retry цикл, начнем с новой моделью
                    else:
                        raise Exception(f"GPT модель {gpt_model} не найдена и нет альтернатив")
                else:
                    continue  # Retry для других попыток
                
            except Exception as e:
                # Другие ошибки
                tbox_log(f"⚠️ GPT ошибка на попытке {attempt+1}/3: {str(e)}", META, "WARN", conf)
                if attempt < 2:
                    delay = 8 + (2 * attempt)  # 8, 10, 12 секунд
                    tbox_log(f"⏳ Пауза {delay} сек перед повтором...", META, "INFO", conf)
                    time.sleep(delay)
                    continue
                else:
                    raise Exception(f"GPT перевод завершился с ошибкой после 3 попыток: {str(e)}")

# Smart chunking functions
def minimal_smart_chunking(text, max_chars=10000, buffer_size=150):
    """
    Умное разбиение текста на чанки с учетом структуры и препинаний
    
    Алгоритм работы:
    1. Ищет границы предложений (.!?) в приоритете
    2. Если нет предложений, ищет границы абзацев (\n\n)
    3. Если нет абзацев, ищет границы строк (\n)
    4. В крайнем случае разбивает по максимальному размеру
    
    Args:
        text (str): Исходный текст для разбиения
        max_chars (int): Максимальный размер чанка в символах
        buffer_size (int): Размер буфера для поиска границ
        
    Returns:
        list: Список чанков текста
    """
    chunks = []
    pos = 0
    
    while pos < len(text):
        # Если осталось меньше max_chars, берем все остальное
        if len(text) - pos <= max_chars:
            chunks.append(text[pos:])
            break
        
        # Ищем умную границу в пределах max_chars + buffer_size
        end_pos = find_smart_boundary(text, pos, max_chars, buffer_size)
        chunk = text[pos:end_pos]
        chunks.append(chunk)
        pos = end_pos
    
    return chunks

def find_smart_boundary(text, start_pos, max_chars, buffer_size):
    """
    Умный поиск границы для разбиения текста с учетом структуры
    
    Приоритеты поиска границ:
    1. Конец предложения (.!? + пробел/новая строка)
    2. Конец абзаца (\n\n)
    3. Конец строки (\n)
    4. Максимальный размер (строго не превышая лимит)
    
    Args:
        text (str): Полный текст
        start_pos (int): Начальная позиция для поиска
        max_chars (int): Максимальный размер чанка (строгий лимит)
        buffer_size (int): Размер буфера для поиска
        
    Returns:
        int: Позиция границы для разбиения
    """
    # Вычисляем область поиска - ищем в диапазоне [max_chars - buffer_size, max_chars]
    # Это позволяет найти границу максимально близко к лимиту, но не превышая его
    search_start = start_pos + max_chars - buffer_size
    search_end = min(len(text), start_pos + max_chars)
    
    # Если начало поиска выходит за пределы текста, ищем от start_pos
    if search_start < start_pos:
        search_start = start_pos
    
    search_text = text[search_start:search_end]
    
    # 1. Ищем конец предложения (.!? с последующим пробелом или новой строкой)
    for i in range(len(search_text) - 1, -1, -1):
        if search_text[i] in '.!?':
            # Проверяем, что после точки есть пробел или новая строка
            if i + 1 < len(search_text) and search_text[i + 1] in ' \n':
                actual_pos = search_start + i + 1
                chunk_size = actual_pos - start_pos
                min_size = max_chars * 0.3  # Минимум 30% от максимума
                # Возвращаем только если чанк не превышает лимит
                if chunk_size >= min_size and chunk_size <= max_chars:
                    return actual_pos
    
    # 2. Ищем конец абзаца (\n\n)
    for i in range(len(search_text) - 1, -1, -1):
        if i > 0 and search_text[i-1] == '\n' and search_text[i] == '\n':
            actual_pos = search_start + i
            chunk_size = actual_pos - start_pos
            min_size = max_chars * 0.3
            if chunk_size >= min_size and chunk_size <= max_chars:
                return actual_pos
    
    # 3. Ищем конец строки (\n)
    for i in range(len(search_text) - 1, -1, -1):
        if search_text[i] == '\n':
            actual_pos = search_start + i
            chunk_size = actual_pos - start_pos
            min_size = max_chars * 0.3
            if chunk_size >= min_size and chunk_size <= max_chars:
                return actual_pos
    
    # 4. Ищем запятую с пробелом (менее предпочтительно)
    for i in range(len(search_text) - 1, -1, -1):
        if search_text[i] == ',' and i + 1 < len(search_text) and search_text[i + 1] == ' ':
            actual_pos = search_start + i + 1
            chunk_size = actual_pos - start_pos
            min_size = max_chars * 0.3
            if chunk_size >= min_size and chunk_size <= max_chars:
                return actual_pos
    
    # 5. В крайнем случае разбиваем точно по максимальному размеру
    end_pos = min(len(text), start_pos + max_chars)
    return end_pos

def translate_md_chunk(chunk, part_num, total_parts, author, conf, prompts, prompt_code, include_original, use_gpt=False, use_gpt_last=False):
    """Перевод MD-текста с retry-логикой и переключением модели"""
    global CURRENT_MODEL, gemini_fallback_list, current_gemini_index
    
    if use_gpt:
        # Перевод через GPT с fallback
        try:
            # Сохраняем флаг use_gpt_last в conf для использования в translate_with_gpt
            if use_gpt_last:
                conf['use_gpt_last'] = True
        
            return translate_with_gpt(chunk, part_num, total_parts, author, conf, prompts, prompt_code, include_original)
        except Exception as e:
            tbox_log(f"GPT перевод не удался: {str(e)}", META, "ERROR", conf)
            response = input(f"❗️ GPT ошибка: {str(e)}\nПереключиться на Gemini? (да/нет): ")
            if response.lower() in ['да', 'yes', 'y']:
                tbox_log("Переключение на Gemini по запросу пользователя", META, "INFO", conf)
                return translate_md_chunk(chunk, part_num, total_parts, author, conf, prompts, prompt_code, include_original, use_gpt=False)
            else:
                raise Exception(f"Перевод отменен пользователем после GPT ошибки: {str(e)}")
    
    # Стандартный перевод через Gemini
    api_key = conf.get('API_KEY', '').split('#')[0].strip()
    
    # Получить промпт из справочника
    base_prompt = prompts.get(prompt_code, prompts.get('GENERIC', 'Переведи на русский: {chunk}'))
    
    # Подставить плейсхолдеры
    prompt = base_prompt.format(
        author=author,
        part_num=part_num,
        total_parts=total_parts,
        chunk=chunk
    )
    
    # Если include_original, добавить инструкцию о цитатах
    if include_original:
        prompt += "\n\nВКЛЮЧАЙ ОРИГИНАЛЬНЫЙ ТЕКСТ ЦИТАТ В СКОБКАХ ПОСЛЕ ПЕРЕВОДА."

    while True:
        # Маркер начала чанка с текущей моделью
        header_marker = f"\n\n--- [PART {part_num}/{total_parts} | MODEL: {CURRENT_MODEL}] ---\n"
        
        # 3 попытки для текущей модели
        for attempt in range(3):
            try:
                api_ver = get_api_ver(CURRENT_MODEL)
                url = f"https://generativelanguage.googleapis.com/{api_ver}/models/{CURRENT_MODEL}:generateContent?key={api_key}"
                
                tbox_log(f"Попытка {attempt+1}/3: {CURRENT_MODEL}", META, "INFO", conf)
                
                r = requests.post(url, json={"contents": [{"parts": [{"text": prompt}]}]}, timeout=120)
                
                if r.status_code == 200:
                    # Безопасное извлечение содержимого
                    try:
                        response_data = r.json()
                        if (response_data and 
                            'candidates' in response_data and 
                            len(response_data['candidates']) > 0 and
                            'content' in response_data['candidates'][0] and
                            'parts' in response_data['candidates'][0]['content'] and
                            len(response_data['candidates'][0]['content']['parts']) > 0 and
                            'text' in response_data['candidates'][0]['content']['parts'][0]):
                            
                            content = response_data['candidates'][0]['content']['parts'][0]['text']
                            tbox_log(f"✅ Часть {part_num}/{total_parts}", META, "INFO", conf)
                            return header_marker + content
                        else:
                            tbox_log(f"❗️ Пустой или некорректный ответ от API", META, "ERROR", conf)
                            # Продолжаем со следующей попыткой или модели
                    except (KeyError, IndexError, TypeError) as e:
                        tbox_log(f"❗️ Ошибка парсинга ответа: {str(e)[:30]}", META, "ERROR", conf)
                        # Продолжаем со следующей попыткой или модели
                
                # Дифференцированная обработка ошибок
                elif r.status_code == 404:
                    # Модель не существует/удалена — менять сразу
                    tbox_log(f"⚠️ Модель {CURRENT_MODEL} не найдена (404)", META, "WARN", conf)
                    
                    new_model = find_working_model(api_key, conf, exclude_model=CURRENT_MODEL)
                    if new_model and new_model != CURRENT_MODEL:
                        CURRENT_MODEL = new_model
                        tbox_log(f"Переключились на: {CURRENT_MODEL}", META, "INFO", conf)
                        break
                    else:
                        return f"\n\n[❗️ ОШИБКА 404: Модель не найдена]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                
                elif r.status_code == 400:
                    # Bad Request — ошибка в запросе/промпте, не в модели
                    tbox_log(f"❗️ Ошибка 400 Bad Request: проблема с промптом", META, "ERROR", conf)
                    return f"\n\n[❗️ ОШИБКА 400: Bad Request]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                
                elif r.status_code == 429:
                    # Дифференцированная обработка 429 ошибок
                    error_msg = r.text.lower() if r.text else ""
                    
                    if 'requests per minute' in error_msg or 'rpm' in error_msg:
                        # RPM лимит - retry с задержкой
                        tbox_log(f"⚠️ 429 RPM лимит на попытке {attempt+1}/3", META, "WARN", conf)
                        if attempt < 2:
                            delay = 15 + (5 * attempt)  # 15, 20, 25 секунд
                            tbox_log(f"⏳ Пауза {delay} сек для RPM", META, "INFO", conf)
                            time.sleep(delay)
                            continue
                        else:
                            return f"\n\n[❗️ ОШИБКА 429 RPM: Лимит запросов в минуту]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                    
                    elif 'requests per day' in error_msg or 'rpd' in error_msg or 'daily' in error_msg:
                        # RPD лимит - смена модели немедленно
                        tbox_log(f"🔄 429 RPD лимит - смена модели", META, "WARN", conf)
                        
                        new_model = find_working_model(api_key, conf, exclude_model=CURRENT_MODEL)
                        if new_model and new_model != CURRENT_MODEL:
                            CURRENT_MODEL = new_model
                            tbox_log(f"Переключились на: {CURRENT_MODEL}", META, "INFO", conf)
                            break
                        else:
                            return f"\n\n[❗️ ОШИБКА 429 RPD: Все модели исчерпаны]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                    
                    else:
                        # Неизвестный тип 429 - дифференцированная обработка
                        if part_num > 1 and attempt < 2:
                            # Не первый чанк - retry с паузой
                            tbox_log(f"⚠️ 429 неизвестный тип на попытке {attempt+1}/3 - retry", META, "WARN", conf)
                            delay = 10 + (3 * attempt)  # 10, 13, 16 секунд
                            tbox_log(f"⏳ Пауза {delay} сек перед повтором...", META, "INFO", conf)
                            time.sleep(delay)
                            continue
                        else:
                            # Первый чанк или последняя попытка - смена модели
                            tbox_log(f"🔄 429 неизвестный тип - смена модели", META, "WARN", conf)
                            
                            new_model = find_working_model(api_key, conf, exclude_model=CURRENT_MODEL)
                            if new_model and new_model != CURRENT_MODEL:
                                CURRENT_MODEL = new_model
                                tbox_log(f"Переключились на: {CURRENT_MODEL}", META, "INFO", conf)
                                break
                            else:
                                return f"\n\n[❗️ ОШИБКА 429: Нет доступных моделей]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                
                elif r.status_code in [500, 503]:
                    # Server Error, Service Unavailable — повторять
                    tbox_log(f"⚠️ Ошибка {r.status_code} на попытке {attempt+1}/3", META, "WARN", conf)
                    
                    if attempt == 2:  # На последней попытке ищем замену
                        new_model = find_working_model(api_key, conf)
                        if new_model and new_model != CURRENT_MODEL:
                            CURRENT_MODEL = new_model
                            tbox_log(f"Переключились на: {CURRENT_MODEL}", META, "INFO", conf)
                            break
                        else:
                            return f"\n\n[❗️ ОШИБКА {r.status_code}: Все модели недоступны]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                
                else:
                    # Неизвестная ошибка — обрабатывать как временная
                    tbox_log(f"⚠️ API Error {r.status_code} на попытке {attempt+1}/3", META, "WARN", conf)
                    
                    if attempt == 2:  # На последней попытке ищем замену
                        new_model = find_working_model(api_key, conf)
                        if new_model and new_model != CURRENT_MODEL:
                            CURRENT_MODEL = new_model
                            tbox_log(f"Переключились на: {CURRENT_MODEL}", META, "INFO", conf)
                            break
                        else:
                            return f"\n\n[❗️ ОШИБКА {r.status_code}: Все модели недоступны]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                    
            except Exception as e:
                tbox_log(f"❗️ Сетевая ошибка: {str(e)[:50]}", META, "ERROR", conf)
            
            if attempt < 2:
                time.sleep(15)  # Пауза перед повтором
        else:
            # Цикл завершился без break — все 3 попытки провалились и не было смены модели
            return f"\n\n[❗️ КРИТИЧЕСКАЯ ОШИБКА ЧАНКА {part_num}]\n[ИСХОДНИК]:\n{chunk[:300]}..."

def render_md_to_docx(md_text, doc):
    """Отрисовка Markdown в параграфы Word"""
    for line in md_text.split('\n'):
        line = line.strip()
        if not line: continue
        p = doc.add_paragraph()
        if line.startswith('#'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.lstrip('#').strip())
            run.bold = True; run.font.size = Pt(14)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                run = p.add_run(part.replace('**', ''))
                run.font.name = 'Times New Roman'; run.font.size = Pt(12)
                if part.startswith('**'): run.bold = True

def main():
    global CURRENT_MODEL, PROGRAM_SHOULD_EXIT
    
    # Если флаг установлен, выходим сразу
    if PROGRAM_SHOULD_EXIT:
        sys.exit(0)
    
    # Парсинг аргументов
    parser = argparse.ArgumentParser(description='TBox Translator')
    parser.add_argument('file', nargs='?', help='Путь к файлу или имя файла')
    parser.add_argument('-txt', action='store_true', help='Использовать последний txt из TXT_RAW')
    parser.add_argument('-md', action='store_true', help='Использовать последний md из MD_DIR')
    parser.add_argument('-s', action='store_true', help='Включать оригинал цитат')
    parser.add_argument('-p', '--prompt', type=str, help='Код промпта (TORAH, FICTION, GENERIC)', dest='prompt')
    parser.add_argument('--gpt', action='store_true', help='Использовать GPT модель из конфига')
    parser.add_argument('--gptlast', action='store_true', help='Использовать последнюю доступную GPT модель')
    parser.add_argument('--max-chars', type=int, help='Максимальное количество символов в чанке (перезаписывает GPT_MAX_CHARS)')
    parser.add_argument('--no-archive', action='store_true', help='Не перемещать обработанные файлы в архив (для тренировок)')
    args = parser.parse_args()
    
    CONF = load_tbox_config()
    if not CONF: 
        print("❗️ Ошибка: Конфиг не найден")
        return
    
    # Если указан --max-chars, обновляем конфиг
    if args.max_chars:
        CONF['GPT_MAX_CHARS'] = str(args.max_chars)
        tbox_log(f"📏 Установлен лимит чанка: {args.max_chars} символов", META, "INFO", CONF)
    
    # Проверка GPT конфигурации если требуется
    use_gpt = False
    use_gpt_last = False
    
    if args.gpt:
        tbox_log("Запрошен перевод через GPT (модель из конфига)", META, "INFO", CONF)
        if check_gpt_config(CONF):
            use_gpt = True
            CONF['use_gpt_last'] = False  # Явно указываем что НЕ ищем модель
            tbox_log("GPT конфигурация корректна, используем GPT из конфига", META, "INFO", CONF)
        else:
            print("❗️ Ошибка: GPT конфигурация некорректна")
            response = input("Переключиться на Gemini? (да/нет): ")
            if response.lower() in ['да', 'yes', 'y']:
                tbox_log("Переключение на Gemini по запросу пользователя", META, "INFO", CONF)
                use_gpt = False
            else:
                tbox_log("Перевод отменен пользователем", META, "INFO", CONF)
                return
    elif args.gptlast:
        tbox_log("Запрошен перевод через последнюю GPT модель", META, "INFO", CONF)
        if check_gpt_config(CONF):
            use_gpt = True
            use_gpt_last = True
            tbox_log("GPT конфигурация корректна, ищем последнюю модель", META, "INFO", CONF)
        else:
            print("❗️ Ошибка: GPT конфигурация некорректна")
            response = input("Переключиться на Gemini? (да/нет): ")
            if response.lower() in ['да', 'yes', 'y']:
                tbox_log("Переключение на Gemini по запросу пользователя", META, "INFO", CONF)
                use_gpt = False
            else:
                tbox_log("Перевод отменен пользователем", META, "INFO", CONF)
                return
    else:
        pass #tbox_log("Используем Gemini (по умолчанию)", META, "INFO", CONF)
    
    # Загрузка промптов
    PROMPTS = load_prompts(CONF)
    
    # Выбор промпта
    prompt_code = args.prompt or CONF.get('DEFAULT_PROMPT', 'TORAH')
    if prompt_code not in PROMPTS:
        prompt_code = 'GENERIC'
    tbox_log(f"👉 Выбран промпт: {prompt_code}", META, "INFO", CONF)
    
    # Инициализация глобальной модели
    CURRENT_MODEL = CONF.get('MODEL_GEMINI', 'gemini-2.0-flash').strip()
    
    # Пути из конфига
    txt_raw_dir = CONF.get('TXT_RAW')           # 02_TXT/raw
    md_dir = CONF.get('MD_DIR', '02_TXT/MD')    # Для md файлов
    out_dir = CONF.get('DOC_TRANSLATED')        # 03_DOC/TRANSLATED
    arh_dir = CONF.get('ARH_TXT')               # 05_ARH/TXT

    # --- ВЫБОР ФАЙЛА ---
    target = None
    if args.file:
        if os.path.exists(args.file):
            target = args.file
        elif args.file.endswith('.txt'):
            potential_path = os.path.join(txt_raw_dir, args.file)
            if os.path.exists(potential_path):
                target = potential_path
        elif args.file.endswith('.md'):
            potential_path = os.path.join(md_dir, args.file)
            if os.path.exists(potential_path):
                target = potential_path
    
    if not target:
        if args.md:
            # Последний md из MD_DIR
            files = glob.glob(os.path.join(md_dir, "*.md"))
            if files:
                target = max(files, key=os.path.getmtime)
        elif args.txt or not (args.md or args.txt):
            # Последний txt из TXT_RAW (по умолчанию)
            files = glob.glob(os.path.join(txt_raw_dir, "*.txt"))
            if files:
                target = max(files, key=os.path.getmtime)
    
    if not target:
        tbox_log("Нет подходящих файлов для перевода", META, "INFO", CONF)
        return

    file_name = os.path.basename(target)
    author = file_name.split('-')[1].replace('_', ' ') if '-' in file_name else "Раввин"
    
    tbox_log(f"🟢 СТАРТ ПЕРЕВОДА: {file_name}", META, "START", CONF)

    with open(target, 'r', encoding='utf-8') as f:
        full_text = f.read()

    # Выводим общую информацию о тексте
    total_chars = len(full_text)
    
    # Определяем модель и её лимиты
    if use_gpt:
        gpt_model = CONF.get('MODEL_GPT', 'gpt-4o')
        chunk_limits = get_gpt_chunk_limits(gpt_model)
        
        # Проверяем лимит из конфига
        gpt_max_chars = int(CONF.get('GPT_MAX_CHARS', '0'))
        
        # Используем минимальный лимит: из конфига или из модели
        if gpt_max_chars > 0:
            max_chars = min(chunk_limits['recommended_chunk_size'], gpt_max_chars)
            tbox_log(f"👉 Будет использована GPT модель: {gpt_model}", META, "INFO", CONF)
            tbox_log(f"ℹ️  Лимиты модели (токенов): {chunk_limits['input_limit']:,} входных / {chunk_limits['output_limit']:,} выходных", META, "INFO", CONF)
            tbox_log(f"ℹ️  Лимит чанка: {max_chars:,} символов (минимум из {chunk_limits['recommended_chunk_size']:,} и {gpt_max_chars:,})", META, "INFO", CONF)
        else:
            max_chars = chunk_limits['recommended_chunk_size']
            tbox_log(f"👉 Будет использована GPT модель: {gpt_model}", META, "INFO", CONF)
            tbox_log(f"ℹ️  Лимиты модели (токенов): {chunk_limits['input_limit']:,} входных / {chunk_limits['output_limit']:,} выходных", META, "INFO", CONF)
            tbox_log(f"ℹ️  Лимит чанка: {max_chars:,} символов", META, "INFO", CONF)
    else:
        # Для Gemini также определяем лимиты динамически
        gemini_model = CURRENT_MODEL if 'CURRENT_MODEL' in globals() else 'gemini-1.5-flash'
        chunk_limits = get_gemini_chunk_limits(gemini_model)
        tbox_log(f"👉 Будет использована Gemini модель: {gemini_model}", META, "INFO", CONF)
        tbox_log(f"ℹ️  Лимиты модели (токенов): {chunk_limits['input_limit']:,} входных / {chunk_limits['output_limit']:,} выходных", META, "INFO", CONF)
        tbox_log(f"ℹ️  Лимит чанка: {chunk_limits['recommended_chunk_size']:,} символов", META, "INFO", CONF)
        max_chars = chunk_limits['recommended_chunk_size']
    
    # Smart chunking с учетом препинаний
    chunks = minimal_smart_chunking(full_text, max_chars=max_chars, buffer_size=150)
    
    tbox_log(f"ℹ️  В тексте: {total_chars:,} символов, разбит на {len(chunks)} чанков", META, "INFO", CONF)
    
    doc = Document()
    # Настройка страницы
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    doc.add_heading(f"Перевод лекции: {author}", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, chunk in enumerate(chunks, 1):
        # Прогресс-бар: сколько символов взято и сколько осталось
        chars_processed = sum(len(c) for c in chunks[:i])
        chars_remaining = sum(len(c) for c in chunks[i:])

        tbox_log("=" * 60, META, "INFO", CONF)
        tbox_log(f"✂️  Чанк {i}/{len(chunks)}: {len(chunk):,} символов, осталось; {chars_remaining:,}", META, "INFO", CONF)
        
        translated_md = translate_md_chunk(chunk, i, len(chunks), author, CONF, PROMPTS, prompt_code, args.s, use_gpt, use_gpt_last)
        render_md_to_docx(translated_md, doc)
        if i < len(chunks):
            time.sleep(12)

    # Сохранение
    os.makedirs(out_dir, exist_ok=True)
    res_name = f"TR_{file_name.replace('.txt', '.docx').replace('.md', '.docx')}"
    res_path = os.path.join(out_dir, res_name)
    
    # Исправляем compatibility mode для современного Word (убираем Compatible Mode)
    try:
        from docx.oxml import parse_xml
        from docx.oxml.ns import qn
        
        # Получаем settings.xml
        settings = doc.settings._element
        
        # Находим или создаем w:compat
        compat = settings.find(qn('w:compat'))
        if compat is None:
            from lxml import etree
            compat = etree.SubElement(settings, qn('w:compat'))
        
        # Удаляем старые настройки compatibilityMode
        for setting in compat.findall(qn('w:compatSetting')):
            if setting.get(qn('w:name')) == 'compatibilityMode':
                compat.remove(setting)
        
        # Добавляем современный compatibilityMode (16 = Word 2016+)
        from lxml import etree
        new_setting = etree.SubElement(compat, qn('w:compatSetting'))
        new_setting.set(qn('w:name'), 'compatibilityMode')
        new_setting.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
        new_setting.set(qn('w:val'), '16')  # Word 2016+
        
    except Exception as e:
        tbox_log(f"❗️ Ошибка исправления compatibility mode: {e}", META, "WARNING", CONF)
    
    doc.save(res_path)
    
    # Архивируем (если не указан --no-archive)
    if not args.no_archive:
        os.makedirs(arh_dir, exist_ok=True)
        try:
            shutil.move(target, os.path.join(arh_dir, f"TR_DONE_{file_name}"))
            tbox_log(f"🟢 ГОТОВО: {res_name} (файл перемещен в архив)", META, "DONE", CONF)
        except Exception as e:
            tbox_log(f"Файл сохранен, архивирование ошибка: {e}", META, "WARNING", CONF)
    else:
        tbox_log(f"🟢 ГОТОВО: {res_name} (файл оставлен на месте, --no-archive)", META, "DONE", CONF)
    
    # Выводим итоговую статистику по GPT токенам
    if use_gpt:
        print_gpt_token_stats(CONF)

if __name__ == "__main__":
    main()
    
    # Проверяем флаг завершения после формирования списка
    if PROGRAM_SHOULD_EXIT:
        sys.exit(0)