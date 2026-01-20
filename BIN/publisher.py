import os, re
from docx import Document
from datetime import datetime

# v1.3-stable - 2026-01-19
# Цепочка: DOC2PUB -> FINAL
# Задача: Финальная чистка текста и удаление технических заголовков

DRAFT_DIR = os.path.expanduser("~/Documents/AI_Lab/DOC2PUB")
FINAL_DIR = os.path.expanduser("~/Documents/AI_Lab/FINAL")

def get_now():
    return datetime.now().strftime("%H:%M:%S")

def clean_paragraph_text(text):
    """Удаляет мусор, который иногда вставляет нейросеть между кусками"""
    # Удаляем строки типа "Часть 1 из 5", "Продолжение следует" и т.д.
    patterns = [
        r"Часть \d+ из \d+",
        r"Часть \d+",
        r"--- PAGE BREAK ---",
        r"Продолжение лекции",
        r"Перевод части \d+"
    ]
    for pattern in patterns:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE)
    return text.strip()

def process_file(file_path):
    file_name = os.path.basename(file_path)
    print(f"[{get_now()}] Обработка: {file_name}...")
    
    doc = Document(file_path)
    new_doc = Document()
    
    # Копируем заголовок (если он есть)
    if doc.paragraphs:
        title = doc.paragraphs[0].text
        new_doc.add_heading(title, 0)

    # Проходим по всем параграфам, начиная со второго
    for para in doc.paragraphs[1:]:
        cleaned_text = clean_paragraph_text(para.text)
        
        if cleaned_text:
            new_p = new_doc.add_paragraph()
            # Переносим форматирование (жирный/курсив)
            for run in para.runs:
                clean_run_text = clean_paragraph_text(run.text)
                if clean_run_text:
                    new_run = new_p.add_run(clean_run_text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic

    # Формируем имя финального файла
    # Убираем "Draft_" или "Source_" для красоты, если хочешь, 
    # но лучше оставить Source_, чтобы отличать двуязычные файлы.
    final_name = file_name.replace("Draft_", "READY_").replace("Source_", "FINAL_SOURCE_")
    save_path = os.path.join(FINAL_DIR, final_name)
    
    new_doc.save(save_path)
    return final_name

def main():
    if not os.path.exists(DRAFT_DIR):
        print(f"[{get_now()}] Ошибка: Папка {DRAFT_DIR} не найдена.")
        return
    
    os.makedirs(FINAL_DIR, exist_ok=True)
    
    # Берем все .docx файлы из папки черновиков
    files = [f for f in os.listdir(DRAFT_DIR) if f.endswith(".docx") and not f.startswith("~$")]
    
    if not files:
        print(f"[{get_now()}] В папке DOC2PUB нет файлов для финализации.")
        return

    print(f"[{get_now()}] --- ЗАПУСК ПАБЛИШЕРА (Файлов: {len(files)}) ---")
    
    for f in files:
        full_path = os.path.join(DRAFT_DIR, f)
        result = process_file(full_path)
        print(f"[{get_now()}] Создан чистовик: {result}")

    print(f"[{get_now()}] --- ГОТОВО ---")

if __name__ == "__main__":
    main()