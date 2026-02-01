def tbox_save_to_docx(text, file_path, title="Lecture"):
    """
    Умная верстка DOCX: RTL/LTR, Justify, Markdown (Bold/Italic)
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document()
    
    # 1. Настройка страницы
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # 2. Определение направления (RTL/LTR)
    # Проверка на наличие иврита или арабского
    is_rtl = bool(re.search(r'[\u0590-\u05FF\u0600-\u06FF]', text[:1000]))
    
    # Заголовок документа
    header = doc.add_heading(title, 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. Обработка блоков текста
    for block in text.split('\n'):
        block = block.strip()
        if not block: continue
        
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        
        # Установка направления и выравнивания
        if is_rtl:
            fmt.rtl = True
            fmt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # А. Обработка Заголовков (#)
        if block.startswith('#'):
            clean_text = block.lstrip('#').strip()
            run = p.add_run(clean_text)
            run.bold = True
            run.font.size = Pt(14)
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Б. Парсинг Markdown (**жирный**, *курсив*)
            # Регулярка ловит блоки: ***жир-курс***, **жирный**, *курсив*
            parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', block)
            for part in parts:
                if not part: continue
                
                if part.startswith('***') and part.endswith('***'):
                    run = p.add_run(part[3:-3])
                    run.bold = True
                    run.italic = True
                elif part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    run = p.add_run(part)
                
                run.font.name = 'Arial' if is_rtl else 'Times New Roman'
                run.font.size = Pt(12)

    doc.save(file_path)