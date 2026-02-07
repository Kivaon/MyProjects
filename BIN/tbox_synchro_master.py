import os, sys, re, json
from datetime import datetime
from difflib import SequenceMatcher
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement, qn
import tbox_utils as utils

# --- ПАСПОРТ ---
VERSION = "v2.0.universal"
DATE    = "2026-02-01"
NAME    = os.path.basename(__file__)
META = {"name": NAME, "version": VERSION, "date": DATE}

# --- КОНСТАНТЫ ---
SEGMENT_SIZE = 600  # Размер сегмента для "стены текста"
MIN_MATCH_RATIO = 0.15  # Уменьшил с 0.3 до 0.15 для реальных файлов
PARTIAL_MATCH_LENGTH = 50  # Минимальная длина для частичного совпадения

class TextSegment:
    """Класс для представления текстового сегмента"""
    def __init__(self, text, original_index, is_header=False):
        self.text = text.strip()
        self.original_index = original_index
        self.is_header = is_header
        self.aligned_with = None
        
    def __repr__(self):
        preview = self.text[:50].replace('\n', ' ') + "..." if len(self.text) > 50 else self.text
        return f"Segment[{self.original_index}]: {preview}"

class SynchroMaster:
    """Основной класс синхронизации"""
    
    def __init__(self, conf):
        self.conf = conf
        self.left_segments = []
        self.right_segments = []
        self.alignment_map = []
        
    def log(self, message, level="INFO"):
        """Логирование через tbox_utils"""
        utils.tbox_log(message, META, level, self.conf)
        
    def preprocess_text(self, text):
        """Предварительная обработка текста"""
        # Виртуальное сегментирование для "стены текста"
        if text.count('\n') < len(text) / 1000:  # Если мало переносов строк
            return self.segment_wall_text(text)
        else:
            return self.segment_structured_text(text)
    
    def segment_wall_text(self, text):
        """Разбиваю стену текста на сегменты"""
        segments = []
        sentences = re.split(r'([.!?]+)', text)
        
        current_segment = ""
        for i in range(0, len(sentences), 2):
            if i + 1 < len(sentences):
                sentence = sentences[i] + sentences[i + 1]
            else:
                sentence = sentences[i]
            
            if len(current_segment + sentence) < SEGMENT_SIZE:
                current_segment += sentence
            else:
                if current_segment.strip():
                    segments.append(current_segment.strip())
                current_segment = sentence
        
        if current_segment.strip():
            segments.append(current_segment.strip())
            
        return segments
    
    def segment_structured_text(self, text):
        """Разбиваю структурированный текст"""
        segments = []
        lines = text.split('\n')
        current_segment = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_segment.strip():
                    segments.append(current_segment.strip())
                    current_segment = ""
                continue
                
            # Заголовки всегда отдельные сегменты
            if line.startswith('#'):
                if current_segment.strip():
                    segments.append(current_segment.strip())
                segments.append(line)
                current_segment = ""
            else:
                if current_segment:
                    current_segment += " " + line
                else:
                    current_segment = line
        
        if current_segment.strip():
            segments.append(current_segment.strip())
            
        return segments
    
    def create_segments(self, segments_list):
        """Создаю объекты TextSegment"""
        result = []
        for i, text in enumerate(segments_list):
            is_header = text.strip().startswith('#')
            result.append(TextSegment(text, i, is_header))
        return result
    
    def calculate_text_similarity(self, text1, text2):
        """Улучшенное сравнение текстов с частичными совпадениями"""
        if not text1 or not text2:
            return 0.0
        
        # Очистка для сравнения
        clean1 = re.sub(r'\s+', ' ', text1.lower().strip())
        clean2 = re.sub(r'\s+', ' ', text2.lower().strip())
        
        # Полное совпадение
        full_ratio = SequenceMatcher(None, clean1, clean2).ratio()
        
        # Если полное совпадение хорошее - возвращаем его
        if full_ratio >= 0.6:
            return full_ratio
        
        # Ищем частичные совпадения
        # Разбиваем на слова и ищем общие последовательности
        words1 = clean1.split()
        words2 = clean2.split()
        
        if len(words1) < 3 or len(words2) < 3:
            return full_ratio  # Слишком короткие для частичного анализа
        
        # Ищем самую длинную общую подпоследовательность слов
        matcher = SequenceMatcher(None, words1, words2)
        match = matcher.find_longest_match(0, len(words1), 0, len(words2))
        
        if match.size >= 3:  # Хотя бы 3 слова подряд
            # Длина совпадения относительно меньшего текста
            partial_ratio = match.size / min(len(words1), len(words2))
            
            # Учитываем и полное, и частичное совпадение
            combined_ratio = max(full_ratio, partial_ratio * 0.8)  # Частичное с весом 0.8
            
            # Добавляем бонус за длинное совпадение
            if match.size >= 5:
                combined_ratio += 0.1
            if match.size >= 8:
                combined_ratio += 0.1
                
            return min(combined_ratio, 1.0)
        
        return full_ratio
    
    def find_anchors(self):
        """Нахожу якоря для выравнивания - от производного к основному"""
        self.log("Поиск якорей для выравнивания...")
        
        # Определяем кто производный (меньший файл)
        left_headers = [s for s in self.left_segments if s.is_header]
        right_headers = [s for s in self.right_segments if s.is_header]
        
        # Выбираем производный файл (меньшее количество заголовков)
        if len(left_headers) <= len(right_headers):
            primary_headers, secondary_headers = left_headers, right_headers
            primary_is_left = True
        else:
            primary_headers, secondary_headers = right_headers, left_headers
            primary_is_left = False
        
        self.log(f"Производный файл: {'левый' if primary_is_left else 'правый'} ({len(primary_headers)} заголовков)")
        self.log(f"Основной файл: {'правый' if primary_is_left else 'левый'} ({len(secondary_headers)} заголовков)")
        
        # Для каждого заголовка производного ищем лучший матч в основном
        for primary_header in primary_headers:
            best_match = None
            best_ratio = 0
            
            for secondary_header in secondary_headers:
                if secondary_header.aligned_with:
                    continue
                    
                # Улучшенное сравнение
                ratio = self.calculate_text_similarity(primary_header.text, secondary_header.text)
                
                if ratio > best_ratio and ratio > MIN_MATCH_RATIO:
                    best_ratio = ratio
                    best_match = secondary_header
            
            if best_match:
                primary_header.aligned_with = best_match
                best_match.aligned_with = primary_header
                self.log(f"Якорь найден: {primary_header.text[:30]}... ↔ {best_match.text[:30]}... (сходство: {best_ratio:.2f})")
    
    def align_remaining_segments(self):
        """Выравниваю оставшиеся сегменты - от производного к основному"""
        self.log("Выравнивание оставшихся сегментов...")
        
        # Определяем кто производный (меньший файл)
        if len(self.left_segments) <= len(self.right_segments):
            primary_segments, secondary_segments = self.left_segments, self.right_segments
            primary_is_left = True
        else:
            primary_segments, secondary_segments = self.right_segments, self.left_segments
            primary_is_left = False
        
        self.log(f"Производный файл: {'левый' if primary_is_left else 'правый'} ({len(primary_segments)} сегментов)")
        self.log(f"Основной файл: {'правый' if primary_is_left else 'левый'} ({len(secondary_segments)} сегментов)")
        
        # Собираем уже выравненные заголовки
        aligned_primary = set()
        aligned_secondary = set()
        
        for seg in primary_segments:
            if seg.aligned_with:
                aligned_primary.add(seg)
                aligned_secondary.add(seg.aligned_with)
        
        # Выравниваем остальные сегменты
        primary_unaligned = [s for s in primary_segments if s not in aligned_primary]
        secondary_unaligned = [s for s in secondary_segments if s not in aligned_secondary]
        
        # Для каждого сегмента производного ищем лучший матч в основном
        for primary_seg in primary_unaligned:
            best_match = None
            best_ratio = 0
            
            for secondary_seg in secondary_unaligned:
                # Улучшенное сравнение
                ratio = self.calculate_text_similarity(primary_seg.text, secondary_seg.text)
                
                if ratio > best_ratio and ratio > MIN_MATCH_RATIO:
                    best_ratio = ratio
                    best_match = secondary_seg
            
            if best_match:
                # Нашли пару
                primary_seg.aligned_with = best_match
                best_match.aligned_with = primary_seg
                secondary_unaligned.remove(best_match)
            else:
                # Не нашли пару - оставляем пустым
                primary_seg.aligned_with = None
        
        # Строим карту выравнивания в правильном порядке
        self.build_alignment_map(primary_is_left)
        
        # Добавляем оставшиеся невыровненные сегменты основного файла
        remaining_secondary = [s for s in secondary_segments if s not in aligned_secondary and not s.aligned_with]
        for secondary_seg in remaining_secondary:
            if primary_is_left:
                self.alignment_map.append((None, secondary_seg))
            else:
                self.alignment_map.append((secondary_seg, None))
    
    def build_alignment_map(self, primary_is_left):
        """Строим карту выравнивания в правильном порядке"""
        # Очищаем карту
        self.alignment_map = []
        
        if primary_is_left:
            # Левый - производный, правый - основной
            for left_seg in self.left_segments:
                if left_seg.aligned_with:
                    self.alignment_map.append((left_seg, left_seg.aligned_with))
                else:
                    self.alignment_map.append((left_seg, None))
        else:
            # Правый - производный, левый - основной
            for right_seg in self.right_segments:
                if right_seg.aligned_with:
                    self.alignment_map.append((right_seg.aligned_with, right_seg))
                else:
                    self.alignment_map.append((None, right_seg))
    
    def detect_language(self, text):
        """Определяю язык текста с улучшенной детекцией иврита"""
        if not text or not text.strip():
            return 'english'
        
        text_clean = text.strip()
        
        # Иврит имеет приоритет - проверяем первые 100 символов
        hebrew_pattern = r'[\u0590-\u05FF\uFB1D-\uFB4F]'  # Включая буквы с диакритикой
        if re.search(hebrew_pattern, text_clean[:100]):
            return 'hebrew'
        
        # Русский
        if re.search(r'[а-яёА-ЯЁ]', text_clean[:100]):
            return 'russian'
        
        # По умолчанию английский
        return 'english'
    
    def clean_text_for_xml(self, text):
        """Очистка текста от недопустимых XML символов с сохранением иврита"""
        if not text:
            return ""
        
        # Удаляем NULL bytes и control characters, но сохраняем иврит
        cleaned = ""
        for char in text:
            char_code = ord(char)
            if char_code == 0:  # NULL
                continue
            elif char_code < 32 and char_code not in [9, 10, 13]:  # control chars except tab, LF, CR
                continue
            elif char_code > 0x10FFFF:  # beyond Unicode range
                continue
            else:
                cleaned += char
        
        return cleaned
    
    def format_cell(self, cell, text):
        """Форматирую ячейку с macOS-специфической поддержкой иврита"""
        if not text.strip():
            return
            
        # Очищаем текст от недопустимых символов
        clean_text = self.clean_text_for_xml(text)
        language = self.detect_language(clean_text)
        
        # Отладка
        if language == 'hebrew':
            self.log(f"Иврит detected: {clean_text[:30]}...")
        
        # Очищаем ячейку полностью
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)
        
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        
        if language == 'hebrew':
            # macOS-специфический метод для иврита
            try:
                # Метод 1: macOS шрифты и XML
                from docx.oxml import parse_xml
                from docx.oxml.ns import qn
                
                # Создаем RTL параграф с macOS-специфическими настройками
                paragraph._element.set(qn('w:bidi'), '1')
                paragraph._element.set(qn('w:jc'), 'right')  # выравнивание вправо
                paragraph._element.set(qn('w:textDirection'), 'rtl')
                
                # macOS шрифты для иврита
                hebrew_fonts = ['David', 'New Peninim', 'Arial', 'Times New Roman']
                
                # Добавляем текст через XML с macOS шрифтами
                run_xml = f'''
                <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:rPr>
                        <w:rtl w:val="1"/>
                        <w:rFonts w:ascii="{hebrew_fonts[0]}" w:hAnsi="{hebrew_fonts[0]}" w:cs="{hebrew_fonts[0]}"/>
                        <w:sz w:val="22"/>
                        <w:lang w:val="he-IL"/>
                    </w:rPr>
                    <w:t xml:space="preserve">{clean_text}</w:t>
                </w:r>
                '''
                
                run_element = parse_xml(run_xml)
                paragraph._element.append(run_element)
                
                self.log(f"Иврит добавлен через macOS XML метод со шрифтом {hebrew_fonts[0]}")
                
            except Exception as e:
                self.log(f"macOS XML метод не сработал: {e}, пробую стандартный")
                
                # Метод 2: Стандартный с macOS шрифтами
                run = paragraph.add_run(clean_text)
                paragraph.paragraph_format.right_to_left = True
                
                # macOS шрифты для иврита
                mac_hebrew_fonts = ['David', 'New Peninim', 'Arial Hebrew', 'Arial']
                
                # Пробуем разные шрифты
                font_set = False
                for font in mac_hebrew_fonts:
                    try:
                        run.font.name = font
                        run.font.rtl = True
                        run.font.size = Pt(11)
                        font_set = True
                        self.log(f"Использован macOS шрифт: {font}")
                        break
                    except:
                        continue
                
                if not font_set:
                    # Fallback на Arial
                    run.font.name = 'Arial'
                    run.font.rtl = True
                    run.font.size = Pt(11)
                
                # macOS-специфические BIDI настройки
                try:
                    paragraph._element.set(qn('w:bidi'), '1')
                    paragraph._element.set(qn('w:textDirection'), 'rtl')
                    run._element.set(qn('w:rtl'), '1')
                    run._element.set(qn('w:lang'), 'he-IL')
                except:
                    pass
                
                try:
                    from docx.oxml.shared import OxmlElement
                    bidi = OxmlElement('w:bidi')
                    bidi.set(qn('w:val'), '1')
                    run._rPr.append(bidi)
                    
                    # Добавляем язык для macOS
                    lang = OxmlElement('w:lang')
                    lang.set(qn('w:val'), 'he-IL')
                    run._rPr.append(lang)
                    
                except:
                    pass
                
        else:
            # Стандартный метод для русского/английского
            run = paragraph.add_run(clean_text)
            paragraph.paragraph_format.right_to_left = False
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            
            try:
                paragraph._element.set(qn('w:bidi'), '0')
                run._element.set(qn('w:rtl'), '0')
            except:
                pass
    
    def create_docx(self, left_path, right_path, output_path):
        """Создаю DOCX документ с улучшенной поддержкой иврита"""
        output_full = os.path.abspath(output_path)
        self.log(f"Создание DOCX: {os.path.basename(output_path)}")
        self.log(f"Полный путь: {output_full}")
        
        doc = Document()
        
        # Альбомная ориентация
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        
        # Минимальные поля
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        
        # Заголовок
        title = doc.add_heading(f"Сверка: {os.path.basename(left_path)} ↔ {os.path.basename(right_path)}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация
        info_para = doc.add_paragraph()
        info_para.add_run(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        info_para.add_run(f"Версия: TBox Synchro Master {VERSION}")
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Пробел
        
        # Создаем таблицу
        table = doc.add_table(rows=len(self.alignment_map) + 1, cols=2)
        table.style = 'Table Grid'
        
        # Заголовки таблицы
        header_cells = table.rows[0].cells
        header_cells[0].text = os.path.basename(left_path)
        header_cells[1].text = os.path.basename(right_path)
        
        # Заполняем таблицу
        for row_idx, (left_seg, right_seg) in enumerate(self.alignment_map, 1):
            cells = table.rows[row_idx].cells
            
            # Левая ячейка
            if left_seg:
                self.format_cell(cells[0], left_seg.text)
            else:
                cells[0].text = ""
                self.format_cell(cells[0], "")
            
            # Правая ячейка
            if right_seg:
                self.format_cell(cells[1], right_seg.text)
            else:
                cells[1].text = ""
                self.format_cell(cells[1], "")
        
        doc.save(output_path)
        self.log(f"DOCX сохранен: {os.path.basename(output_path)}", "DONE")
        self.log(f"Полный путь: {output_full}", "DONE")
    
    def read_file_safe(self, file_path):
        """Безопасное чтение файла с разными кодировками и отладкой иврита"""
        encodings = ['utf-8', 'utf-8-sig', 'cp1255', 'cp1252', 'latin-1', 'utf-16']
        
        self.log(f"Попытка чтения файла: {os.path.basename(file_path)}")
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                    
                # Проверяем наличие иврита
                if re.search(r'[\u0590-\u05FF]', text):
                    self.log(f"Найден иврит в файле, кодировка: {encoding}")
                    # Нормализация Unicode для иврита
                    import unicodedata
                    text = unicodedata.normalize('NFC', text)
                    self.log(f"Иврит нормализован, пример: {text[:50]}...")
                
                return text
                
            except UnicodeDecodeError:
                continue
            except Exception as e:
                self.log(f"Ошибка чтения файла {file_path}: {e}", "ERROR")
                return None
        
        self.log(f"Не удалось прочитать файл {file_path} ни в одной кодировке", "ERROR")
        return None
    
    def calculate_similarity_stats(self):
        """Подсчитываю статистику сходства с улучшенным алгоритмом"""
        total_pairs = len(self.alignment_map)
        aligned_pairs = 0
        total_similarity = 0
        high_similarity = 0
        medium_similarity = 0
        low_similarity = 0
        
        for left_seg, right_seg in self.alignment_map:
            if left_seg and right_seg:
                aligned_pairs += 1
                
                # Используем улучшенное сравнение
                ratio = self.calculate_text_similarity(left_seg.text, right_seg.text)
                total_similarity += ratio
                
                # Категоризируем сходство
                if ratio >= 0.8:
                    high_similarity += 1
                elif ratio >= 0.5:
                    medium_similarity += 1
                else:
                    low_similarity += 1
        
        # Процентное соотношение
        alignment_rate = (aligned_pairs / total_pairs * 100) if total_pairs > 0 else 0
        avg_similarity = (total_similarity / aligned_pairs * 100) if aligned_pairs > 0 else 0
        
        stats = {
            'total_pairs': total_pairs,
            'aligned_pairs': aligned_pairs,
            'empty_left': total_pairs - aligned_pairs,
            'empty_right': total_pairs - aligned_pairs,
            'alignment_rate': alignment_rate,
            'avg_similarity': avg_similarity,
            'high_similarity': high_similarity,
            'medium_similarity': medium_similarity,
            'low_similarity': low_similarity
        }
        
        return stats
    
    def log_similarity_stats(self, stats):
        """Вывожу статистику в лог"""
        self.log("\n📊 СТАТИСТИКА СХОДСТВА:")
        self.log(f"  Всего пар:           {stats['total_pairs']}")
        self.log(f"  Выровненных пар:      {stats['aligned_pairs']} ({stats['alignment_rate']:.1f}%)")
        self.log(f"  Пустых ячеек слева:  {stats['empty_left']}")
        self.log(f"  Пустых ячеек справа: {stats['empty_right']}")
        self.log(f"  Среднее сходство:    {stats['avg_similarity']:.1f}%")
        self.log(f"  Высокое сходство:     {stats['high_similarity']} (≥80%)")
        self.log(f"  Среднее сходство:     {stats['medium_similarity']} (50-79%)")
        self.log(f"  Низкое сходство:      {stats['low_similarity']} (<50%)")
        
        if stats['alignment_rate'] < 50:
            self.log("⚠️  Низкий процент выравнивания - возможно, файлы не соответствуют друг другу", "WARN")
        elif stats['avg_similarity'] < 60:
            self.log("⚠️  Низкое сходство текстов - возможно, разные версии или язык", "WARN")
        else:
            self.log("✅ Хорошее выравнивание и сходство текстов", "DONE")
    
    def process_files(self, left_path, right_path, output_path):
        """Основной процесс синхронизации"""
        left_full = os.path.abspath(left_path)
        right_full = os.path.abspath(right_path)
        output_full = os.path.abspath(output_path)
        
        self.log(f"Старт синхронизации:")
        self.log(f"  Левый файл:  {left_full}")
        self.log(f"  Правый файл: {right_full}")
        self.log(f"  Выходной:    {output_full}")
        
        # Читаем файлы с обработкой кодировок
        left_text = self.read_file_safe(left_path)
        right_text = self.read_file_safe(right_path)
        
        if left_text is None or right_text is None:
            self.log("Ошибка чтения файлов", "ERROR")
            return
        
        # Предварительная обработка
        self.log("Предварительная обработка текстов...")
        left_segments_raw = self.preprocess_text(left_text)
        right_segments_raw = self.preprocess_text(right_text)
        
        # Создаем сегменты
        self.left_segments = self.create_segments(left_segments_raw)
        self.right_segments = self.create_segments(right_segments_raw)
        
        self.log(f"Сегментов создано: левый={len(self.left_segments)}, правый={len(self.right_segments)}")
        
        # Выравнивание
        self.find_anchors()
        self.align_remaining_segments()
        
        self.log(f"Выравнивание завершено: {len(self.alignment_map)} пар")
        
        # Считаем и выводим статистику сходства
        stats = self.calculate_similarity_stats()
        self.log_similarity_stats(stats)
        
        # Создаем DOCX
        self.create_docx(left_path, right_path, output_path)
        
        # Спросить об открытии файла
        try:
            response = input(f"\nОткрыть созданный файл? {output_full} [y/N]: ").strip().lower()
            if response in ['y', 'yes', 'да', 'д']:
                import subprocess
                import platform
                
                if platform.system() == 'Darwin':  # macOS
                    subprocess.run(['open', output_full])
                elif platform.system() == 'Windows':
                    subprocess.run(['start', output_full], shell=True)
                elif platform.system() == 'Linux':
                    subprocess.run(['xdg-open', output_full])
                
                self.log("Файл открыт для просмотра")
        except KeyboardInterrupt:
            self.log("Отмена открытия файла")
        except Exception as e:
            self.log(f"Ошибка открытия файла: {e}", "WARN")
    
    def export_corrections(self, docx_path, output_md_path):
        """Экспорт правок из правой колонки"""
        self.log(f"Экспорт правок из: {os.path.basename(docx_path)}")
        
        doc = Document(docx_path)
        extracted_text = []
        
        # Пропускаем заголовок таблицы
        for row in doc.tables[0].rows[1:]:
            right_cell = row.cells[1]
            if right_cell.text.strip():
                extracted_text.append(right_cell.text.strip())
        
        # Сохраняем как MD
        with open(output_md_path, 'w', encoding='utf-8') as f:
            f.write('\n\n'.join(extracted_text))
        
        self.log(f"Правки экспортированы: {os.path.basename(output_md_path)}", "DONE")

def find_file_pairs(conf, mode):
    """Поиск пар файлов для разных режимов"""
    base_dir = conf.get('BASE_DIR', os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    
    if mode == "refinery":
        raw_dir = conf.get('TXT_RAW', '02_TXT/raw')
        md_dir = conf.get('MD_DIR', '02_TXT/MD')
        
        # Ищем последнюю пару raw ↔ md по соответствию имен
        raw_files = {f: os.path.join(raw_dir, f) for f in os.listdir(raw_dir) if f.endswith('.txt')}
        md_files = {f: os.path.join(md_dir, f) for f in os.listdir(md_dir) if f.endswith('.md')}
        
        # Ищем пары с похожими именами
        best_pair = None
        best_time = 0
        
        for raw_name, raw_path in raw_files.items():
            # Убираем расширения и префиксы для сравнения
            raw_base = re.sub(r'_raw\.txt$', '', raw_name)
            raw_base = re.sub(r'^\d{6}_\d{6}_', '', raw_base)  # убираем таймстемп
            
            for md_name, md_path in md_files.items():
                md_base = re.sub(r'\.md$', '', md_name)
                md_base = re.sub(r'^\d{6}_\d{6}_', '', md_base)  # убираем таймстемп
                
                # Проверяем сходство имен
                if raw_base.lower() in md_name.lower() or md_name.lower() in raw_name.lower():
                    # Нашли пару - проверяем время
                    pair_time = max(os.path.getmtime(raw_path), os.path.getmtime(md_path))
                    if pair_time > best_time:
                        best_time = pair_time
                        best_pair = (raw_path, md_path)
        
        return best_pair if best_pair else (None, None)
    
    elif mode == "translate":
        md_dir = conf.get('MD_DIR', '02_TXT/MD')
        doc_dir = conf.get('DOC_TRANSLATED', '03_DOC/TRANSLATED')
        
        # Ищем последнюю пару md ↔ docx по соответствию имен
        md_files = {f: os.path.join(md_dir, f) for f in os.listdir(md_dir) if f.endswith('.md')}
        doc_files = {f: os.path.join(doc_dir, f) for f in os.listdir(doc_dir) if f.endswith('.docx')}
        
        # Ищем пары с похожими именами
        best_pair = None
        best_time = 0
        
        for md_name, md_path in md_files.items():
            # Убираем расширения и префиксы
            md_base = re.sub(r'\.md$', '', md_name)
            md_base = re.sub(r'^\d{6}_\d{6}_', '', md_base)
            
            for doc_name, doc_path in doc_files.items():
                doc_base = re.sub(r'\.docx$', '', doc_name)
                doc_base = re.sub(r'^Ready_', '', doc_base)  # убираем префикс Ready_
                doc_base = re.sub(r'^\d{6}_\d{6}_', '', doc_base)
                
                # Проверяем сходство имен
                if md_base.lower() in doc_name.lower() or doc_name.lower() in md_name.lower():
                    # Нашли пару - проверяем время
                    pair_time = max(os.path.getmtime(md_path), os.path.getmtime(doc_path))
                    if pair_time > best_time:
                        best_time = pair_time
                        best_pair = (md_path, doc_path)
        
        return best_pair if best_pair else (None, None)
    
    return None, None

def main():
    """Главная функция"""
    # Загрузка конфигурации
    conf = utils.load_local_config()
    if not conf:
        print("Ошибка: config.txt не найден")
        return
    
    synchro = SynchroMaster(conf)
    
    # Парсинг аргументов
    if len(sys.argv) < 2:
        print("TBox Synchro Master v2.0")
        print("\nИспользование:")
        print("  python tbox_synchro_master.py manual <left_file> <right_file>")
        print("  python tbox_synchro_master.py autoref")
        print("  python tbox_synchro_master.py autotr")
        print("  python tbox_synchro_master.py export <docx_file> <output_md>")
        return
    
    command = sys.argv[1].lower()
    
    if command == "manual":
        if len(sys.argv) < 4:
            print("Ошибка: нужны два файла для сравнения")
            return
        
        left_path = sys.argv[2]
        right_path = sys.argv[3]
        
        if not os.path.exists(left_path) or not os.path.exists(right_path):
            print("Ошибка: один из файлов не найден")
            return
        
        # Создаем выходное имя
        timestamp = datetime.now().strftime("%y%m%d_%H%M%S")
        left_name = os.path.splitext(os.path.basename(left_path))[0]
        right_name = os.path.splitext(os.path.basename(right_path))[0]
        output_path = f"synchro_{timestamp}_{left_name}_vs_{right_name}.docx"
        
        synchro.process_files(left_path, right_path, output_path)
    
    elif command == "autoref":
        left_path, right_path = find_file_pairs(conf, "refinery")
        if not left_path or not right_path:
            synchro.log("Пара файлов для refinery не найдена", "ERROR")
            return
        
        timestamp = datetime.now().strftime("%y%m%d_%H%M%S")
        output_path = f"synchro_refinery_{timestamp}.docx"
        
        synchro.process_files(left_path, right_path, output_path)
    
    elif command == "autotr":
        left_path, right_path = find_file_pairs(conf, "translate")
        if not left_path or not right_path:
            synchro.log("Пара файлов для translate не найдена", "ERROR")
            return
        
        timestamp = datetime.now().strftime("%y%m%d_%H%M%S")
        output_path = f"synchro_translate_{timestamp}.docx"
        
        synchro.process_files(left_path, right_path, output_path)
    
    elif command == "export":
        if len(sys.argv) < 4:
            print("Ошибка: нужны docx файл и выходной md файл")
            return
        
        docx_path = sys.argv[2]
        output_md = sys.argv[3]
        
        if not os.path.exists(docx_path):
            print("Ошибка: DOCX файл не найден")
            return
        
        synchro.export_corrections(docx_path, output_md)
    
    else:
        print(f"Неизвестная команда: {command}")

if __name__ == "__main__":
    main()
