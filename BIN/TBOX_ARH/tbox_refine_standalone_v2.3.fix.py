import sys, os, requests, re, time
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- ПАСПОРТ (MANIFEST) ---
VERSION = "v2.3.fix"
DATE    = "2026-01-24"
NAME    = os.path.basename(__file__)
META    = {"name": NAME, "version": VERSION}

def tbox_log(msg, meta, level="INFO", conf=None):
    print(f"[{datetime.now().strftime('%H:%M:%S')}] [{meta['name']}] [{level}] {msg}")

def load_local_config():
    """Локальная копия загрузчика конфига, чтобы не зависеть от внешних файлов"""
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    config_path = os.path.join(base_dir, "config.txt")
    conf = {}
    if not os.path.exists(config_path): return None
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.split('#')[0].strip()
                if '=' in line:
                    k, v = line.split('=', 1)
                    conf[k.strip()] = v.strip()
        actual_base = conf.get('BASE_DIR', base_dir)
        for k in conf:
            if '${BASE_DIR}' in conf[k]:
                conf[k] = conf[k].replace('${BASE_DIR}', actual_base)
        return conf
    except: return None

def find_best_model(api_key, is_large_text=True):
    """Ищет Gemini и игнорирует слабую Gemma для больших текстов"""
    for ver in ["v1beta", "v1"]:
        url = f"https://generativelanguage.googleapis.com/{ver}/models?key={api_key}"
        try:
            r = requests.get(url, timeout=5)
            if r.status_code != 200: continue
            models = r.json().get('models', [])
            
            # Фильтруем: только Gemini (у них лимиты выше)
            candidates = []
            for m in models:
                m_name = m['name'].replace('models/', '')
                if 'generateContent' not in m.get('supportedGenerationMethods', []): continue
                if is_large_text and "gemma" in m_name.lower(): continue
                candidates.append(m_name)
            
            for m_name in sorted(candidates, reverse=True):
                test_url = f"https://generativelanguage.googleapis.com/{ver}/models/{m_name}:generateContent?key={api_key}"
                try:
                    res = requests.post(test_url, json={"contents": [{"parts": [{"text": "hi"}]}]}, timeout=3)
                    if res.status_code == 200:
                        return m_name, ver
                except: continue
        except: continue
    return None, None

def tbox_chunk_text(text, max_chars=10000):
    chunks = []
    current_chunk = []
    current_length = 0
    for line in text.split('\n'):
        if current_length + len(line) > max_chars:
            chunks.append("\n".join(current_chunk))
            current_chunk = [line]; current_length = len(line)
        else:
            current_chunk.append(line); current_length += len(line)
    if current_chunk: chunks.append("\n".join(current_chunk))
    return chunks
def tbox_save_to_docx(text, file_path, title="Lecture"):
    """
    Умная верстка DOCX: RTL/LTR, Justify, Markdown (Bold/Italic)
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    doc = Document()
    
    # 1. Настройка страницы
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # 2. Определение направления (RTL/LTR)
    # Проверка на наличие иврита или арабского
    is_rtl = bool(re.search(r'[\u0590-\u05FF\u0600-\u06FF]', text[:1000]))
    
    # Заголовок документа
    header = doc.add_heading(title, 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. Обработка блоков текста
    for block in text.split('\n'):
        block = block.strip()
        if not block: continue
        
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        
        # УСТАНОВКА ВЫРАВНИВАНИЯ "ПО ШИРИНЕ" ДЛЯ ВСЕХ ЯЗЫКОВ
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Это сделает ровными ОБА края
        
        if is_rtl:
            fmt.rtl = True
            # Для иврита JUSTIFY работает корректно вместе с rtl=True
        else:
            fmt.rtl = False

        # А. Обработка Заголовков (#)
        if block.startswith('#'):
            clean_text = block.lstrip('#').strip()
            run = p.add_run(clean_text)
            run.bold = True
            run.font.size = Pt(14)
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Б. Парсинг Markdown (**жирный**, *курсив*)
            # Регулярка ловит блоки: ***жир-курс***, **жирный**, *курсив*
            parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', block)
            for part in parts:
                if not part: continue
                
                if part.startswith('***') and part.endswith('***'):
                    run = p.add_run(part[3:-3])
                    run.bold = True
                    run.italic = True
                elif part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    run = p.add_run(part)
                
                run.font.name = 'Arial' if is_rtl else 'Times New Roman'
                run.font.size = Pt(12)

    doc.save(file_path)

def main():
    if len(sys.argv) < 2: return print("Usage: ref <file>")
    raw_path = sys.argv[1]
    
    CONF = load_local_config()
    if not CONF: return print("Config not found!")
    
    with open(raw_path, 'r', encoding='utf-8') as f:
        content = f.read()

    api_key = CONF.get('API_KEY', '').split('#')[0].strip()
    
    # Решаем, нужна ли нам мощная модель
    is_large = len(content) > 5000
    m_name, m_ver = find_best_model(api_key, is_large_text=is_large)

    if not m_name:
        tbox_log("Нет подходящих моделей Gemini онлайн.", META, "ERROR")
        return

    chunks = tbox_chunk_text(content, max_chars=10000)
    tbox_log(f"Текст: {len(content)} симв. ({len(chunks)} ч.). Модель: {m_name}", META, "INFO")
    
    refined_full = ""
    for i, chunk in enumerate(chunks, 1):
        # Добавляем паузу ПЕРЕД запросом, если это не первая часть
        if i > 1:
            tbox_log("Пауза 10 секунд для обхода лимита RPM...", META, "INFO")
            time.sleep(10)
        tbox_log(f"Обработка {i}/{len(chunks)}...", META, "START")
        prompt = (
            "ТЫ — МАСТЕР ВЕРСТКИ И ПУНКТУАЦИИ. СЛОВА НЕ МЕНЯТЬ.\n"
            "ТВОЯ ЗАДАЧА: Оформить авторский текст, сохранив каждое слово в оригинальном виде.\n\n"
            "ПРАВИЛА ОФОРМЛЕНИЯ:\n"
            "1. ПУНКТУАЦИЯ: Если в RAW-тексте не хватает знаков препинания или заглавных букв — расставь их, не меняя порядок слов.\n"
            "2. АБЗАЦЫ: Разделяй текст на абзацы СТРОГО по смыслу. Не делай их слишком мелкими или огромными.\n"
            "3. ЗАГОЛОВКИ: Если в тексте нет заголовков, то можно дать название темы для логических блоков, оформи его строкой с символом '#' в начале.\n"
            "4. ЦИТАТЫ: Все приводимые автором цитаты, выдержки или ссылки на источники выделяй жирным шрифтом '**...**'.\n"
            "5. ПОЛНЫЙ ЗАПРЕТ: Никаких комментариев, вступлений, резюме от ИИ или перефразирования. \n"
            "НАЧНИ ВЫВОД СРАЗУ С ПЕРВОГО СЛОВА АВТОРСКОГО ТЕКСТА И ЗАКОНЧИ ПОСЛЕДНИМ\n"
            f"\nТЕКСТ ДЛЯ ВЕРСТКИ:\n{chunk}"
        )
        url = f"https://generativelanguage.googleapis.com/{m_ver}/models/{m_name}:generateContent?key={api_key}"

        try:
            r = requests.post(url, json={"contents": [{"parts": [{"text": prompt}]}]}, timeout=120)
            res = r.json()
            
            if 'candidates' in res and res['candidates'][0].get('content'):
                chunk_result = res['candidates'][0]['content']['parts'][0]['text']
                refined_full += chunk_result.strip() + "\n\n"
            else:
                # Если ИИ заблокировал текст или выдал ошибку
                reason = res.get('error', {}).get('message', 'Safety Filter / Unknown Error')
                tbox_log(f"ЧАСТЬ {i} ОТКЛОНЕНА API: {reason}", META, "ERROR")
                # Добавляем оригинал, чтобы не потерять кусок лекции
                refined_full += f"\n\n[!!! КУСОК НЕ ОБРАБОТАН: {reason}]\n\n" + chunk + "\n\n"
        except Exception as e:
            tbox_log(f"Критический сбой на части {i}: {e}", META, "ERROR")
            refined_full += chunk

    # Сохранение
    base_name = os.path.basename(raw_path).replace("_raw.txt", "").replace(".txt", "")
    
    # Сохраняем Markdown
    md_path = os.path.join(CONF.get('TXT_DIR'), f"{base_name}.md")
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(refined_full)
    
    # Сохраняем DOCX через нашу новую процедуру
    docx_path = os.path.join(CONF.get('DOC_ORIGINALS'), f"{base_name}.docx")
    try:
        tbox_save_to_docx(refined_full, docx_path, title=base_name)
        tbox_log(f"Документы созданы: {base_name}", META, "DONE")
    except Exception as e:
        tbox_log(f"Ошибка сохранения DOCX: {e}", META, "ERROR")

if __name__ == "__main__":
    main()