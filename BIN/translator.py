# --- v4.0 Ultimatum---
import os
import requests
import time
import sys
import warnings
from docx import Document
from datetime import datetime

# Отключаем ворчание SSL
warnings.filterwarnings("ignore")

# --- НАСТРОЙКИ ---
KEY = "AIzaSyB7KMmVo9E9vtvLwX3TnfEybQ8y6qCsQYA"
MODEL = "gemini-2.0-flash"
SOURCE_DIR = os.path.expanduser("~/Documents/AI_Lab/TXT2DOC")
TARGET_DIR = os.path.expanduser("~/Documents/AI_Lab/DOC2PUB")
# -----------------

def get_now():
    return datetime.now().strftime("%H:%M:%S")

def translate_chunk(chunk):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{MODEL}:generateContent?key={KEY}"
    payload = {
        "contents": [{
            "parts": [{
                "text": f"Ты профессиональный редактор. Переведи текст лекции на русский язык ПОЛНОСТЬЮ. Цитаты на иврите оставляй как есть и выделяй жирным (**). Не сокращай текст! ТЕКСТ:\n\n{chunk}"
            }]
        }]
    }
    
    # Пытаемся перевести кусок, пока не получится (цикл упорства)
    while True:
        try:
            r = requests.post(url, json=payload, timeout=60)
            if r.status_code == 200:
                return r.json()['candidates'][0]['content']['parts'][0]['text']
            elif r.status_code == 429:
                print(f"\n[{get_now()}] Лимит Google. Сплю 60 сек...")
                time.sleep(60)
            else:
                print(f"\n[{get_now()}] Ошибка {r.status_code}. Жду 10 сек...")
                time.sleep(10)
        except:
            time.sleep(5)

def main():
    os.makedirs(TARGET_DIR, exist_ok=True)
    all_files = [f for f in os.listdir(SOURCE_DIR) if f.endswith(".txt")]
    if not all_files:
        print("Папка пуста!")
        return
    
    selected_file = max([os.path.join(SOURCE_DIR, f) for f in all_files], key=os.path.getmtime)
    file_name = os.path.basename(selected_file)
    
    print(f"[{get_now()}] >>> СТАРТ ПОЛНОГО ПЕРЕВОДА")
    print(f"[{get_now()}] >>> Файл: {file_name}")
    
    with open(selected_file, "r", encoding="utf-8") as f:
        text_data = f.read()

    # Куски по 5000 знаков — идеально для Gemini
    chunk_size = 5000 
    chunks = [text_data[i:i+chunk_size] for i in range(0, len(text_data), chunk_size)]
    
    doc = Document()
    doc.add_heading(f"Перевод: {file_name}", 0)

    for i, chunk in enumerate(chunks, 1):
        print(f"[{get_now()}] Обработка части {i} из {len(chunks)}...", end="\r")
        
        result = translate_chunk(chunk)
        
        # Добавляем текст в Word с сохранением абзацев
        if result:
            for paragraph in result.split('\n'):
                if paragraph.strip():
                    p = doc.add_paragraph()
                    # Логика для жирного шрифта (**)
                    import re
                    parts = re.split(r'(\*\*.*?\*\*)', paragraph)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            p.add_run(part.replace('**', '')).bold = True
                        else:
                            p.add_run(part)
            
            sys.stdout.write("\033[K")
            print(f"[{get_now()}] ✅ Часть {i} готова.")
            
            # Небольшая пауза, чтобы не частить
            time.sleep(20)

    out_name = f"FULL_TRANS_{file_name.replace('.txt', '.docx')}"
    doc.save(os.path.join(TARGET_DIR, out_name))
    print(f"---")
    print(f"[{get_now()}] 🎉 ФИНАЛ! Файл сохранен: {out_name}")

if __name__ == "__main__":
    main()