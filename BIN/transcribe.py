import os, sys, configparser, re, time, glob, requests
from datetime import datetime
from docx import Document

# --- НАСТРОЙКИ ИЗ CONFIG.TXT ---
def load_config():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    path = os.path.join(base_dir, "config.txt")
    config = configparser.ConfigParser()
    config.optionxform = str
    try:
        with open(path, 'r', encoding='utf-8') as f:
            content = '[DEFAULT]\n' + f.read()
        config.read_string(content)
        conf = config['DEFAULT']
        return conf, base_dir
    except Exception as e:
        print(f"Ошибка чтения конфига: {e}")
        return {}, base_dir

CONF, BASE_DIR = load_config()
API_KEY = CONF.get('API_KEY', '').split('#')[0].strip()
MODEL_NAME = CONF.get('MODEL_GEMINI', 'gemini-2.0-flash').split('#')[0].strip()
LOG_FILE = os.path.join(BASE_DIR, CONF.get('LOG_FILE', 'factory.log').split('#')[0].strip())

# Папки
IN_DIR = os.path.join(BASE_DIR, "02_TXT")
OUT_DIR = os.path.join(BASE_DIR, "03_DOC")

def log(message, level="INFO"):
    t = datetime.now().strftime("%H:%M:%S")
    msg = f"[{t}] [v4.0] [{level}] {message}"
    print(msg)
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(f"[{datetime.now()}] {msg}\n")

def translate_chunk_rest(chunk, part_num, total, author, include_source):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{MODEL_NAME}:generateContent?key={API_KEY}"
    
    source_instr = (
        "Если автор зачитывает источник на иврите: напиши иврит ЖИРНЫМ (**), "
        "а следом перевод на русский КУРСИВОМ В КАВЫЧКАХ (*«...»*)."
    ) if include_source else "НЕ оставляй иврит. Переводи источники сразу на русский ЖИРНЫМ (**)."

    prompt = (
        f"Ты — профессиональный редактор. Переведи часть {part_num}/{total} лекции Рава {author}.\n\n"
        f"ПРАВИЛА:\n1. Обычная речь: только русский.\n2. Цитаты: {source_instr}\n"
        f"3. Контекст: 'Я сказал' или 'Мы учили' — просто переводи.\n4. Тон: связный литературный текст.\n\n"
        f"ТЕКСТ:\n{chunk}"
    )

    payload = {"contents": [{"parts": [{"text": prompt}]}]}
    headers = {"Content-Type": "application/json"}
    
    for attempt in range(5):
        try:
            r = requests.post(url, json=payload, headers=headers, timeout=120)
            
            # --- ВОТ ЗДЕСЬ ПРОВЕРКА ДЕТАЛЕЙ ОШИБКИ 429 ---
            if r.status_code == 429:
                error_details = r.json().get('error', {}).get('message', 'No message')
                wait = 60 * (attempt + 1)
                log(f"Лимит (429). Причина: {error_details}", "WARN")
                log(f"Ожидание {wait} сек (попытка {attempt + 1})...", "INFO")
                time.sleep(wait)
                continue
            # ---------------------------------------------

            r.raise_for_status()
            return r.json()['candidates'][0]['content']['parts'][0]['text']
            
        except requests.exceptions.RequestException as e:
            log(f"Ошибка сетевого запроса: {e}", "ERROR")
            if 'r' in locals() and r.text:
                log(f"Ответ сервера: {r.text[:200]}...", "DEBUG")
            time.sleep(10)
            
    return f"[Ошибка после 5 попыток]"
    
def add_formatted_text(paragraph, text):
    # Умный парсер Markdown ** и *
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part.replace('**', ''))
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part.replace('*', ''))
            run.italic = True
        else:
            paragraph.add_run(part)

def main():
    if not API_KEY: log("API_KEY не найден!", "ERROR"); return
    os.makedirs(OUT_DIR, exist_ok=True)
    include_source = '-s' in sys.argv
    
    # Хелп
    if '-h' in sys.argv:
        print("Использование: tr (последний файл) или tr [часть имени] или tr -s (с ивритом)"); return

    # Поиск файла
    files = glob.glob(os.path.join(IN_DIR, "*.txt"))
    if not files: log("Папка 02_TXT пуста", "ERROR"); return
    
    if len(sys.argv) > 1 and not sys.argv[1].startswith('-'):
        matches = [f for f in files if sys.argv[1] in os.path.basename(f)]
        target = max(matches, key=os.path.getmtime) if matches else max(files, key=os.path.getmtime)
    else:
        target = max(files, key=os.path.getmtime)

    file_name = os.path.basename(target)
    # Извлекаем автора (твой старый метод)
    parts = file_name.split('-')
    author = parts[1].replace("_", " ") if len(parts) > 1 else "Раввин"

    log(f"Старт: {file_name} (Автор: {author})")
    
    with open(target, 'r', encoding='utf-8') as f:
        full_text = f.read()

    # Нарезка по 3500 для стабильности
    chunks = [full_text[i:i+3500] for i in range(0, len(full_text), 3500)]
    doc = Document()
    doc.add_heading(f"Лекция: {author}", 0)

    for i, chunk in enumerate(chunks, 1):
        log(f"Часть {i}/{len(chunks)}...")
        res = translate_chunk_rest(chunk, i, len(chunks), author, include_source)
        
        for line in res.split('\n'):
            if line.strip():
                add_formatted_text(doc.add_paragraph(), line)
        
        time.sleep(12) # Пауза для обхода RPM лимита

    out_file = f"Draft_{file_name.replace('.txt', '.docx')}"
    doc.save(os.path.join(OUT_DIR, out_file))
    log(f"ГОТОВО! Файл в 03_DOC", "DONE")

if __name__ == "__main__":
    main()