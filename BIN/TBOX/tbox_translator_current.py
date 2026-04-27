import os, sys, time, glob, requests, shutil, re, argparse
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from datetime import datetime
from docx import Document   
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tbox_utils import tbox_log

# Попытка импорта OpenAI
try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

# --- MANIFEST ---
# Previous: v6.00.hybrid_pro (2026-01-27) - Hybrid version with smart chunking
VERSION = "v6.12.gpt-simple"
DATE    = "2026-04-27"
NAME    = os.path.basename(__file__)
META    = {"name": NAME, "version": VERSION}

# Глобальные переменные для провайдеров
CURRENT_PROVIDER = "gemini"  # "gemini" или "openai"
CURRENT_MODEL = "gemini-1.5-flash"  # Будет обновлен в init_providers()

def load_tbox_config():
    script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    config_path = os.path.join(script_dir, "_config", "tconfig.txt")
    conf = {}
    if not os.path.exists(config_path): 
        print(f"Ошибка: Конфиг не найден: {config_path}")
        return None
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.split('#')[0].strip()
                if '=' in line:
                    key, val = line.split('=', 1)
                    conf[key.strip()] = val.strip()
        actual_base = conf.get('BASE_DIR', script_dir)
        for key in conf:
            if '${BASE_DIR}' in conf[key]:
                conf[key] = conf[key].replace('${BASE_DIR}', actual_base)
        return conf
    except Exception as e:
        print(f"Ошибка в load_tbox_config: {e}")
        return None

def load_prompts(conf):
    """Загрузка справочника промптов из файла"""
    prompts_dir = conf.get('PROMPTS_DIR', '07_PROMPTS')
    base_dir = conf.get('BASE_DIR', os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    prompts_path = os.path.join(base_dir, prompts_dir, 'prompts.md')
    prompts = {
        'TORAH': "Ты — редактор лекций Рава {author}. Переведи часть {part_num}/{total_parts} на русский.\nСТРОГО СОХРАНЯЙ РАЗМЕТКУ: '#' для заголовков и '**' для выделений.\nСтиль: возвышенный, точный.\nТЕКСТ:\n{chunk}",
        'FICTION': "Переведи художественный текст часть {part_num}/{total_parts} на русский.\nСОХРАНЯЙ стилистику и тон оригинала.\nРазметка: '#' для заголовков, '**' для выделений.\nТЕКСТ:\n{chunk}",
        'GENERIC': "Переведи текст часть {part_num}/{total_parts} на русский.\nТЕКСТ:\n{chunk}"
    }
    if os.path.exists(prompts_path):
        try:
            with open(prompts_path, 'r', encoding='utf-8') as f:
                content = f.read()
            # Парсинг по заголовкам # CODE
            sections = re.split(r'^# (.+)$', content, flags=re.MULTILINE)
            for i in range(1, len(sections), 2):
                code = sections[i].strip()
                text = sections[i+1].strip()
                prompts[code] = text
            tbox_log(f"Загружено промптов: {list(prompts.keys())}", META, "INFO", conf)
        except Exception as e:
            tbox_log(f"Ошибка загрузки промптов: {e}", META, "ERROR", conf)
    else:
        tbox_log(f"Файл промптов не найден, используем встроенные: {prompts_path}", META, "WARN", conf)
    return prompts

def get_api_ver(model_name):
    """Определяет API версию по названию модели"""
    v_match = re.search(r'(\d+)', model_name)
    v_major = int(v_match.group(1)) if v_match else 1
    if v_major >= 2 or any(x in model_name.lower() for x in ['exp', 'beta']):
        return "v1beta"
    return "v1"


def init_providers(conf, use_openai=False):
    """Инициализация провайдеров - Gemini по умолчанию, GPT если указан"""
    global CURRENT_PROVIDER, CURRENT_MODEL
    
    if use_openai:
        # Используем OpenAI GPT
        openai_key = conf.get('OPENAI_API_KEY', '').strip()
        openai_model = conf.get('MODEL_GPT', 'gpt-3.5-turbo').strip()
        
        if not OPENAI_AVAILABLE:
            print("OpenAI не установлен. Установите: pip install openai")
            return False
        
        if not openai_key:
            print("OPENAI_API_KEY не найден в конфигурации")
            return False
        
        CURRENT_PROVIDER = "openai"
        CURRENT_MODEL = openai_model
        print(f"Используем OpenAI GPT: {openai_model}")
        return True
    else:
        # Используем Gemini (по умолчанию)
        gemini_key = conf.get('API_KEY', '').strip()
        gemini_model = conf.get('MODEL_GEMINI', 'gemini-1.5-flash').strip()
        
        if not gemini_key:
            print("API_KEY не найден в конфигурации")
            return False
        
        CURRENT_PROVIDER = "gemini"
        CURRENT_MODEL = gemini_model
        print(f"Используем Gemini: {gemini_model}")
        return True


def translate_with_gemini(prompt, api_key, model):
    """Перевод через Gemini API"""
    api_ver = get_api_ver(model)
    url = f"https://generativelanguage.googleapis.com/{api_ver}/models/{model}:generateContent?key={api_key}"
    
    data = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.3,
            "maxOutputTokens": 8192
        }
    }
    
    response = requests.post(url, json=data, timeout=30)
    return response


def translate_with_openai(prompt, conf):
    """Перевод через OpenAI GPT API"""
    openai_key = conf.get('OPENAI_API_KEY', '').strip()
    openai_model = conf.get('MODEL_GPT', 'gpt-3.5-turbo').strip()
    
    if not OPENAI_AVAILABLE or not openai_key:
        raise Exception("OpenAI GPT недоступен")
    
    client = openai.OpenAI(api_key=openai_key)
    
    response = client.chat.completions.create(
        model=openai_model,
        messages=[
            {"role": "system", "content": "Ты профессиональный переводчик. Переводи текст точно, сохраняя форматирование и смысл."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=8192,
        temperature=0.3
    )
    
    return response

def analyze_429_error(error_text):
    """Анализирует ошибку 429 для определения RPM или RPD"""
    error_lower = error_text.lower()
    
    # Проверка индикаторов RPD (запросов в день)
    rpd_indicators = [
        'per day', 'daily', 'day quota', 'daily quota',
        'requestsperday', 'perday', 'day limit'
    ]
    
    # Проверка индикаторов RPM (запросов в минуту)  
    rpm_indicators = [
        'per minute', 'minute', 'rpm', 'per min',
        'requestsperminute', 'minute quota', 'rate limit'
    ]
    
    if any(indicator in error_lower for indicator in rpd_indicators):
        return 'RPD'  # Дневной лимит превышен - немедленно сменить модель
    elif any(indicator in error_lower for indicator in rpm_indicators):
        return 'RPM'  # Минутный лимит превышен - можно retry с задержкой
    else:
        # По умолчанию RPM для безопасности (можно retry)
        return 'RPM'

def find_working_model(api_key, conf):
    """Automatic selection of working model when main one fails"""
    tbox_log("SEARCHING FOR ALTERNATIVE MODEL...", META, "WARN", conf)
    
    def v_score(name):
        v = re.search(r'(\d+\.?\d*)', name)
        num = float(v.group(1)) if v else 0.0
        score = num + (0.5 if "flash" in name.lower() else 0)
        if any(x in name.lower() for x in ["vision", "pro", "exp"]): score -= 1
        return score

    for api_v in ["v1beta", "v1"]:
        try:
            url = f"https://generativelanguage.googleapis.com/{api_v}/models?key={api_key}"
            r = requests.get(url, timeout=10)
            if r.status_code == 200:
                models = [m['name'].split('/')[-1] for m in r.json().get('models', []) 
                          if 'generateContent' in m.get('supportedGenerationMethods', [])]
                models.sort(key=v_score, reverse=True)
                
                for model in models:
                    test_url = f"https://generativelanguage.googleapis.com/{api_v}/models/{model}:generateContent?key={api_key}"
                    try:
                        tr = requests.post(test_url, json={"contents": [{"parts": [{"text": "test"}]}]}, timeout=5)
                        if tr.status_code == 200:
                            tbox_log(f"Найдена модель: {model}", META, "INFO", conf)
                            return model
                    except:
                        continue
        except:
            continue
    return None

# Smart chunking functions
def minimal_smart_chunking(text, max_chars=10000, buffer_size=150):
    chunks = []
    pos = 0
    
    while pos < len(text):
        chunk_end = min(pos + max_chars, len(text))
        if chunk_end < len(text):
            boundary = find_boundary_minimal(text, pos, chunk_end, buffer_size, max_chars)
            if boundary:
                chunk_end = boundary
        
        chunk_text = text[pos:chunk_end]
        if chunk_text.strip():
            chunks.append(chunk_text)
        pos = chunk_end
    
    return chunks

def find_boundary_minimal(text, start_pos, end_pos, buffer_size, max_chars):
    buffer_start = max(start_pos, end_pos - buffer_size)
    boundary_text = text[buffer_start:end_pos]
    
    for i, char in enumerate(reversed(boundary_text)):
        actual_pos = buffer_start + len(boundary_text) - 1 - i
        if char in '.!?':
            chunk_size = (actual_pos + 1) - start_pos
            min_size = max_chars * 0.3
            if chunk_size >= min_size:
                return actual_pos + 1
    
    for i, char in enumerate(reversed(boundary_text)):
        actual_pos = buffer_start + len(boundary_text) - 1 - i
        if char in ',;:--':
            chunk_size = actual_pos - start_pos
            min_size = max_chars * 0.3
            if chunk_size >= min_size:
                return actual_pos
    
    for i, char in enumerate(reversed(boundary_text)):
        actual_pos = buffer_start + len(boundary_text) - 1 - i
        if char == ' ':
            chunk_size = actual_pos - start_pos
            min_size = max_chars * 0.3
            if chunk_size >= min_size:
                return actual_pos
    
    return end_pos

def translate_md_chunk(chunk, part_num, total_parts, author, conf, prompts, prompt_code, include_original):
    """Перевод MD-текста с retry-логикой и переключением провайдеров"""
    global CURRENT_PROVIDER, CURRENT_MODEL
    
    # Получить промпт из справочника
    base_prompt = prompts.get(prompt_code, prompts.get('GENERIC', 'Переведи на русский: {chunk}'))
    
    # Подставить плейсхолдеры
    prompt = base_prompt.format(
        author=author,
        part_num=part_num,
        total_parts=total_parts,
        chunk=chunk
    )
    
    # Если include_original, добавить инструкцию о цитатах
    if include_original:
        prompt += "\n\nВКЛЮЧАЙ ОРИГИНАЛЬНЫЙ ТЕКСТ ЦИТАТ В СКОБКАХ ПОСЛЕ ПЕРЕВОДА."

    while True:
        # Маркер начала чанка с текущим провайдером
        header_marker = f"\n\n--- [PART {part_num}/{total_parts} | {CURRENT_PROVIDER.upper()}: {CURRENT_MODEL}] ---\n"
        
        # 3 попытки для текущего провайдера
        for attempt in range(3):
            try:
                if CURRENT_PROVIDER == "gemini":
                    api_key = conf.get('API_KEY', '').split('#')[0].strip()
                    response = translate_with_gemini(prompt, api_key, CURRENT_MODEL)
                    
                    if response.status_code == 200:
                        result = response.json()
                        translation = result['candidates'][0]['content']['parts'][0]['text']
                        return header_marker + translation
                    elif response.status_code == 429:
                        error_text = response.text
                        error_type = analyze_429_error(error_text)
                        
                        if error_type == 'RPD':
                            print("Gemini дневной лимит (RPD) - переключение на OpenAI")
                            # Переключаемся на OpenAI
                            if switch_to_openai(conf):
                                break  # Переключились, пробуем заново
                            else:
                                return header_marker + f"\n\n[ОШИБКА: Все провайдеры недоступны]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                        else:
                            # RPM - retry с задержкой
                            if attempt < 2:
                                delay = 10 + (attempt * 5)
                                print(f"Gemini минутный лимит (RPM) - ожидание {delay}s")
                                # Задержка перед retry
                                time.sleep(5)


def switch_to_openai(conf):
    """Переключение на OpenAI GPT"""
    global CURRENT_PROVIDER, CURRENT_MODEL
                                else:
                                    return header_marker + f"\n\n[ОШИБКА: Все провайдеры недоступны]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                    else:
                        raise Exception(f"Gemini API error {response.status_code}: {response.text}")
                
                elif CURRENT_PROVIDER == "openai":
                    response = translate_with_openai(prompt, conf)
                    translation = response.choices[0].message.content
                    return header_marker + translation
                
            except Exception as e:
                error_msg = str(e)
                print(f"Ошибка {CURRENT_PROVIDER.upper()} попытки {attempt+1}: {error_msg}")
                
                if attempt == 2:  # Последняя попытка
                    if CURRENT_PROVIDER == "gemini":
                        # Пробуем переключиться на OpenAI
                        if switch_to_openai(conf):
                            break
                        else:
                            return header_marker + f"\n\n[ОШИБКА: Все провайдеры недоступны]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                    else:
                        return header_marker + f"\n\n[ОШИБКА: OpenAI недоступен]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                
                
def render_md_to_docx(md_text, doc):
    """Отрисовка Markdown в параграфы Word"""
    
    for line in md_text.split('\n'):
        line = line.strip()
        if not line: continue
        p = doc.add_paragraph()
        if line.startswith('#'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.lstrip('#').strip())
            run.bold = True; run.font.size = Pt(14)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                run = p.add_run(part.replace('**', ''))
                run.font.name = 'Times New Roman'; run.font.size = Pt(12)
                if part.startswith('**'): run.bold = True

def switch_to_openai(conf):
    """Переключение на OpenAI GPT"""
    global CURRENT_PROVIDER, CURRENT_MODEL
    
    openai_key = conf.get('OPENAI_API_KEY', '').strip()
    openai_model = conf.get('MODEL_GPT', 'gpt-3.5-turbo').strip()
    
    if not OPENAI_AVAILABLE or not openai_key:
        print("OpenAI GPT недоступен для переключения")
        return False
    
    CURRENT_PROVIDER = "openai"
    CURRENT_MODEL = openai_model
    print(f"Переключение на OpenAI GPT: {openai_model}")
    return True


def main():
    """Основная функция обработки MD-файлов"""
    parser = argparse.ArgumentParser(description='Перевод MD-файлов с мульти-модельной поддержкой')
    parser.add_argument('input_file', help='MD файл для перевода')
    parser.add_argument('--provider', choices=['gemini', 'openai'], help='Принудительный выбор провайдера')
    parser.add_argument('--original', action='store_true', help='Включать оригинальный текст в перевод')
    args = parser.parse_args()
    
    conf = load_tbox_config()
    if not conf:
        print("Ошибка загрузки конфигурации")
        return
    
    # Инициализация провайдеров
    use_openai = args.provider == "openai"
    if not init_providers(conf, use_openai=use_openai):
        print("Ошибка инициализации провайдеров")
        return
    
    # Загрузка промптов
    prompts = load_prompts(conf)
    
    # Обработка файла
    input_path = args.input_file
    if not os.path.exists(input_path):
        print(f"Файл не найден: {input_path}")
        return
    
    print(f"Обработка файла: {input_path}")
    print(f"Провайдер: {CURRENT_PROVIDER.upper()} - {CURRENT_MODEL}")
    
    # Чтение MD файла
    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Разбиение на чанки
    max_chars = 3000
    chunks = []
    start = 0
    
    while start < len(md_content):
        end = start + max_chars
        if end >= len(md_content):
            chunks.append(md_content[start:])
            break
        
        # Ищем ближайший конец предложения/абзаца
        boundary_text = md_content[start:end]
        boundary_pos = find_best_boundary(boundary_text, max_chars)
        
        chunk = md_content[start:start + boundary_pos]
        chunks.append(chunk)
        start += boundary_pos
    
    print(f"Разбито на {len(chunks)} чанков")
    
    # Перевод чанков
    translated_chunks = []
    author = "Автор"  # Можно извлечь из имени файла или конфига
    prompt_code = "GENERIC"  # Можно определить по содержимому
    
    for i, chunk in enumerate(chunks, 1):
        print(f"Перевод чанка {i}/{len(chunks)}...")
        translated = translate_md_chunk(chunk, i, len(chunks), author, conf, prompts, prompt_code, args.original)
        translated_chunks.append(translated)
        
        # Пауза между чанками
        if i < len(chunks):
            time.sleep(2)
    
    # Сохранение результата
    output_path = input_path.replace('.md', '_translated.md')
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(translated_chunks))
    
    print(f"Перевод сохранен: {output_path}")


if __name__ == "__main__":
    main()
            if os.path.exists(potential_path):
                target = potential_path
        elif args.file.endswith('.md'):
            potential_path = os.path.join(md_dir, args.file)
            if os.path.exists(potential_path):
                target = potential_path
    
    if not target:
        if args.md:
            # Последний md из MD_DIR
            files = glob.glob(os.path.join(md_dir, "*.md"))
            if files:
                target = max(files, key=os.path.getmtime)
        elif args.txt or not (args.md or args.txt):
            # Последний txt из TXT_RAW (по умолчанию)
            files = glob.glob(os.path.join(txt_raw_dir, "*.txt"))
            if files:
                target = max(files, key=os.path.getmtime)
    
    if not target:
        tbox_log("Нет подходящих файлов для перевода", META, "INFO", CONF)
        return

    file_name = os.path.basename(target)
    author = file_name.split('-')[1].replace('_', ' ') if '-' in file_name else "Раввин"
    
    tbox_log(f"СТАРТ ПЕРЕВОДА: {file_name}", META, "START", CONF)

    with open(target, 'r', encoding='utf-8') as f:
        full_text = f.read()

    # Smart chunking
    chunks = minimal_smart_chunking(full_text, max_chars=10000, buffer_size=150)
    
    # Создаем Word документ
    doc = Document()
    
    # Настройка страницы
    section = doc.sections[0]
    
    # Размер страницы A4
    from docx.enum.section import WD_ORIENT
    from docx.shared import Mm
    section.page_width = Mm(210)  # A4 width
    section.page_height = Mm(297)  # A4 height
    section.orientation = WD_ORIENT.PORTRAIT
    
    # Поля страницы - Narrow (узкие)
    section.left_margin = Cm(1.27)   # 0.5 inch
    section.right_margin = Cm(1.27)  # 0.5 inch
    section.top_margin = Cm(1.27)    # 0.5 inch
    section.bottom_margin = Cm(1.27) # 0.5 inch
    
    # Нумерация страниц - по центру внизу
    from docx.enum.section import WD_HEADER_FOOTER
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = ""
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    footer_para._element.append(parse_xml(f'<w:fldSimple w:instr="PAGE" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
    
    # Примечание: 2 страницы на листе - это настройка принтера, не документа
    # Для печати 2 страниц на листе нужно настроить в настройках принтера
    doc.add_heading(f"Перевод лекции: {author}", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, chunk in enumerate(chunks, 1):
        tbox_log(f"Перевод чанка {i}/{len(chunks)}...", META, "INFO", CONF)
        translated_md = translate_md_chunk(chunk, i, len(chunks), author, CONF, PROMPTS, prompt_code, args.s)
        render_md_to_docx(translated_md, doc)
        if i < len(chunks):
            time.sleep(12)

    # Сохранение
    os.makedirs(out_dir, exist_ok=True)
    res_name = f"Ready_{file_name.replace('.txt', '.docx').replace('.md', '.docx')}"
    res_path = os.path.join(out_dir, res_name)
    
    # Исправляем compatibility mode для современного Word (убираем Compatible Mode)
    try:
        from docx.oxml import parse_xml
        from docx.oxml.ns import qn
        
        # Получаем settings.xml
        settings = doc.settings._element
        
        # Находим или создаем w:compat
        compat = settings.find(qn('w:compat'))
        if compat is None:
            from lxml import etree
            compat = etree.SubElement(settings, qn('w:compat'))
        
        # Удаляем старые настройки compatibilityMode
        for setting in compat.findall(qn('w:compatSetting')):
            if setting.get(qn('w:name')) == 'compatibilityMode':
                compat.remove(setting)
        
        # Добавляем современный compatibilityMode (16 = Word 2016+)
        from lxml import etree
        new_setting = etree.SubElement(compat, qn('w:compatSetting'))
        new_setting.set(qn('w:name'), 'compatibilityMode')
        new_setting.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
        new_setting.set(qn('w:val'), '16')  # Word 2016+
        
    except Exception as e:
        tbox_log(f"Ошибка исправления compatibility mode: {e}", META, "WARNING", CONF)
    
    doc.save(res_path)
    
    # Архивируем
    os.makedirs(arh_dir, exist_ok=True)
    try:
        shutil.move(target, os.path.join(arh_dir, f"TR_DONE_{file_name}"))
        tbox_log(f"ГОТОВО: {res_name}", META, "DONE", CONF)
    except Exception as e:
        tbox_log(f"Файл сохранен, архивирование ошибка: {e}", META, "WARNING", CONF)

if __name__ == "__main__":
    main()