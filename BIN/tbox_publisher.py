import os, sys, re, shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tbox_utils import tbox_log

# --- MANIFEST ---
VERSION = "v2.9.master_print"
DATE    = "2026-01-23"
NAME    = os.path.basename(__file__)
META    = {"name": NAME, "version": VERSION}

def load_tbox_config():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    base_dir_actual = os.path.dirname(script_dir)
    config_path = os.path.join(base_dir_actual, "config.txt")
    conf = {}
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.split('#')[0].strip()
                if '=' in line:
                    key, val = line.split('=', 1)
                    conf[key.strip()] = val.strip()
        target_base = conf.get('BASE_DIR', base_dir_actual)
        for key in conf:
            if '${BASE_DIR}' in conf[key]:
                conf[key] = conf[key].replace('${BASE_DIR}', target_base)
        return conf
    except: return None

CONF = load_tbox_config()

def clean_markers(text):
    has_error = "[!!! КРИТИЧЕСКАЯ ОШИБКА" in text
    patterns = [r"--- \[PART \d+/\d+ \| MODEL:.*?\] ---", r"\[MODEL:.*?\]", r"Часть \d+ из \d+"]
    clean_text = text
    for p in patterns:
        clean_text = re.sub(p, "", clean_text, flags=re.IGNORECASE)
    return clean_text.strip(), has_error

def process_file(file_path, print_mode="2"):
    file_name = os.path.basename(file_path)
    ts = datetime.now().strftime("%y%m%d_%H%M")
    tbox_log(f"Наводим красоту: {file_name}", META, "INFO", CONF)
    
    try:
        doc = Document(file_path)
        new_doc = Document()
        
        # --- НАСТРОЙКА ГЕОМЕТРИИ ПЕЧАТИ ---
        section = new_doc.sections[0]
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
        if print_mode == "2":
            section.gutter = Cm(0.3) # Зазор в центре листа под сгиб
        
        has_any_error = False

        for para in doc.paragraphs:
            cleaned_text, is_err = clean_markers(para.text)
            if is_err: has_any_error = True
            if not cleaned_text: continue

            # Создаем новый параграф
            new_p = new_doc.add_paragraph()
            
            # 1. Стилистика и Markdown-префиксы
            prefix = ""
            is_heading = False
            if para.style.name.startswith('Heading'):
                level = para.style.name.split()[-1]
                prefix = ("#" * (int(level) + 1) if level.isdigit() else "#") + " "
                is_heading = True
                new_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif "List Bullet" in para.style.name:
                prefix = "* "
            
            # 2. Сборка Markdown внутри строки
            full_md_text = prefix
            for run in para.runs:
                r_text, _ = clean_markers(run.text)
                if not r_text: continue
                if run.bold: r_text = f"**{r_text}**"
                if run.italic: r_text = f"*{r_text}*"
                full_md_text += r_text

            # 3. Наполнение параграфа текстом и форматирование
            new_run = new_p.add_run(full_md_text)
            new_run.font.name = 'Times New Roman'
            new_run.font.size = Pt(11)
            
            # Выравнивание основного текста по обоим краям
            if not is_heading:
                new_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # 4. Нижний колонтитул (номер страницы)
        footer = section.footer.paragraphs[0]
        footer.text = "— PAGE —"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- СОХРАНЕНИЕ НА СКЛАД (04_PUB) ---
        pub_dir = CONF.get('PUB_FINAL')
        os.makedirs(pub_dir, exist_ok=True)
        pub_name = f"FINAL_{ts}_{file_name.replace('Ready_', '').replace('RUS_', '')}"
        new_doc.save(os.path.join(pub_dir, pub_name))

        # --- АРХИВАЦИЯ (05_ARH/DOC) ---
        arh_dir = CONF.get('ARH_DOC')
        os.makedirs(arh_dir, exist_ok=True)
        arh_name = f"PROCESSED_{ts}_{file_name}"
        shutil.move(file_path, os.path.join(arh_dir, arh_name))

        tbox_log(f"ГОТОВО: {pub_name}", META, "DONE", CONF)

    except Exception as e:
        tbox_log(f"Ошибка на файле {file_name}: {e}", META, "ERROR", CONF)

def main():
    if not CONF: return
    in_dir = CONF.get('DOC_TO_PUB')
    
    print(f"\n--- МОНИТОРИНГ ШЛЮЗА ---")
    print(f"Путь: {in_dir}")
    if in_dir and os.path.exists(in_dir):
        files = [f for f in os.listdir(in_dir) if f.lower().endswith(".docx") and not f.startswith("~$")]
        print(f"Найдено документов: {len(files)}")
        for f in files: 
            process_file(os.path.join(in_dir, f), print_mode="2")
    else:
        print("ОШИБКА: Шлюз не найден.")
    print("--- ЗАВЕРШЕНО ---\n")

if __name__ == "__main__":
    main()