# --- ПАРАМЕТРЫ ВЕРСИИ ---
VERSION = "v4.2.prod"
DATE = "2026-01-23"
CHUNK_SIZE = 8000  # Увеличили блок для 2.5/3.0 моделей

import os, sys, configparser, re, time, glob, requests
from datetime import datetime
from docx import Document

# --- ЗАГРУЗКА КОНФИГА ---
def load_config():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    path = os.path.join(base_dir, "config.txt")
    config = configparser.ConfigParser()
    try:
        with open(path, 'r', encoding='utf-8') as f:
            config.read_string('[DEFAULT]\n' + f.read())
        conf = config['DEFAULT']
        return conf, base_dir
    except: return {}, base_dir

CONF, BASE_DIR = load_config()
API_KEY = CONF.get('API_KEY', '').split('#')[0].strip()
MODEL_NAME = CONF.get('MODEL_GEMINI', 'gemini-2.5-flash').strip()
IN_DIR = os.path.join(BASE_DIR, "02_TXT")
OUT_DIR = os.path.join(BASE_DIR, "03_DOC")
LOG_FILE = os.path.join(BASE_DIR, "factory.log")

def log(message, level="INFO"):
    t = datetime.now().strftime("%H:%M:%S")
    msg = f"[{t}] [{VERSION}] [{level}] {message}"
    print(msg)
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(f"{datetime.now().isoformat()} {msg}\n")

def translate_chunk(chunk, part_num, total, author, include_source):
    # Используем ветку v1, так как сканер показал там OK
    url = f"https://generativelanguage.googleapis.com/v1/models/{MODEL_NAME}:generateContent?key={API_KEY}"
    
    instr = (
        "Оставь иврит ЖИРНЫМ (**), перевод КУРСИВОМ (*«...»*)." if include_source 
        else "Переводи источники сразу на русский ЖИРНЫМ (**), иврит удаляй."
    )

    prompt = (
        f"Ты — редактор лекций Рава {author}. Переведи часть {part_num}/{total}.\n"
        f"ПРАВИЛА:\n1. Речь автора: сразу на русский.\n2. Цитаты: {instr}\n"
        f"ТЕКСТ:\n{chunk}"
    )

    payload = {"contents": [{"parts": [{"text": prompt}]}]}
    
    for attempt in range(3):
        try:
            r = requests.post(url, json=payload, timeout=120)
            if r.status_code == 200:
                return r.json()['candidates'][0]['content']['parts'][0]['text']
            if r.status_code == 429:
                log(f"Лимит 429. Ждем {60*(attempt+1)}с...", "WARN")
                time.sleep(60*(attempt+1))
            else:
                log(f"Ошибка {r.status_code}: {r.text[:100]}", "ERROR")
        except Exception as e:
            log(f"Ошибка сети: {e}", "ERROR")
        time.sleep(5)
    return "[Ошибка перевода]"

def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    files = glob.glob(os.path.join(IN_DIR, "*.txt"))
    if not files: log("Нет файлов в 02_TXT", "ERROR"); return

    # Выбор файла (аргумент или последний)
    target = max(files, key=os.path.getmtime)
    if len(sys.argv) > 1 and not sys.argv[1].startswith('-'):
        matches = [f for f in files if sys.argv[1] in f]
        if matches: target = max(matches, key=os.path.getmtime)

    file_name = os.path.basename(target)
    author = file_name.split('-')[1].replace('_', ' ') if '-' in file_name else "Раввин"
    
    log(f"Старт: {file_name} на модели {MODEL_NAME}")
    
    with open(target, 'r', encoding='utf-8') as f:
        text = f.read()

    chunks = [text[i:i+4000] for i in range(0, len(text), 4000)]
    doc = Document()
    doc.add_heading(f"Лекция: {author}", 0)

    for i, chunk in enumerate(chunks, 1):
        log(f"Часть {i}/{len(chunks)}...")
        res = translate_chunk(chunk, i, len(chunks), author, '-s' in sys.argv)
        
        # Простое добавление текста в Word
        p = doc.add_paragraph(res)
        time.sleep(10) # Безопасная пауза

    out_path = os.path.join(OUT_DIR, f"Ready_{file_name.replace('.txt', '.docx')}")
    doc.save(out_path)
    log(f"Завершено: {out_path}", "DONE")

if __name__ == "__main__":
    main()