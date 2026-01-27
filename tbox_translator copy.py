import os, sys, configparser, time, glob, requests, shutil, re
from datetime import datetime
from docx import Document
from tbox_utils import tbox_log

# --- MANIFEST ---
VERSION       = "v4.80.hybrid"
DATE          = "2026-01-23"
NAME          = os.path.basename(__file__)
META          = {"name": NAME, "version": VERSION}

# --- ЕДИНОЕ МЕСТО ДЛЯ ДЕФОЛТОВ ---
DEFAULT_MODEL = "gemini-2.0-flash" 
CHUNK_SIZE    = 10000              

def load_tbox_config():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    base_dir_actual = os.path.dirname(script_dir)
    config_path = os.path.join(base_dir_actual, "config.txt")
    conf = {}
    if not os.path.exists(config_path): return None
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.split('#')[0].strip()
                if '=' in line:
                    key, val = line.split('=', 1)
                    conf[key.strip()] = val.strip()
        target_base = conf.get('BASE_DIR', base_dir_actual)
        for key in conf:
            if '${BASE_DIR}' in conf[key]:
                conf[key] = conf[key].replace('${BASE_DIR}', target_base)
        return conf
    except: return None

# --- ИНИЦИАЛИЗАЦИЯ ---
CONF = load_tbox_config()
API_KEY = CONF.get('API_KEY', '').split('#')[0].strip() if CONF else ""
# Сначала конфиг, если нет - дефолт
CURRENT_MODEL = CONF.get('MODEL_GEMINI', DEFAULT_MODEL).strip() if CONF else DEFAULT_MODEL

def get_api_ver(model_name):
    """Определяет ветку API на основе версии модели"""
    v_match = re.search(r'(\d+)', model_name)
    v_major = int(v_match.group(1)) if v_match else 1
    if v_major >= 2 or any(x in model_name.lower() for x in ['exp', 'beta']):
        return "v1beta"
    return "v1"

def find_working_model():
    """Разведка запускается ТОЛЬКО если основная модель подвела"""
    tbox_log("АВАРИЙНЫЙ ПОИСК: Основная модель недоступна. Ищу замену...", META, "WARN", CONF)
    
    def v_score(name):
        v = re.search(r'(\d+\.?\d*)', name)
        num = float(v.group(1)) if v else 0.0
        score = num + (0.5 if "flash" in name.lower() else 0)
        if any(x in name.lower() for x in ["vision", "pro", "exp"]): score -= 1
        return score

    for ver in ["v1beta", "v1"]:
        url = f"https://generativelanguage.googleapis.com/{ver}/models?key={API_KEY}"
        try:
            r = requests.get(url, timeout=10)
            if r.status_code == 200:
                models = [m['name'].split('/')[-1] for m in r.json().get('models', []) 
                          if 'generateContent' in m.get('supportedGenerationMethods', [])]
                models.sort(key=v_score, reverse=True)
                
                for m in models:
                    test_url = f"https://generativelanguage.googleapis.com/{ver}/models/{m}:generateContent?key={API_KEY}"
                    tr = requests.post(test_url, json={"contents": [{"parts": [{"text": "hi"}]}]}, timeout=5)
                    if tr.status_code == 200:
                        tbox_log(f"Найдена замена: {m}", META, "INFO", CONF)
                        return m
        except: continue
    return None

def translate_chunk(chunk, part_num, total, author, include_source):
    global CURRENT_MODEL
    
    # Маркер начала стыковки для будущего паблишера
    header_marker = f"\n\n--- [PART {part_num}/{total} | MODEL: {CURRENT_MODEL}] ---\n"
    
    instr = ("Оставь иврит ЖИРНЫМ (**), перевод КУРСИВОМ (*«...»*)." if include_source 
             else "Переводи источники сразу на русский ЖИРНЫМ (**), иврит удаляй.")

    prompt = (f"Ты — редактор лекций Рава {author}. Переведи часть {part_num}/{total}.\n"
              f"ПРАВИЛА: 1. Речь автора: русский. 2. Цитаты: {instr}\nТЕКСТ:\n{chunk}")

    for attempt in range(3):
        api_ver = get_api_ver(CURRENT_MODEL)
        url = f"https://generativelanguage.googleapis.com/{api_ver}/models/{CURRENT_MODEL}:generateContent?key={API_KEY}"
        
        try:
            r = requests.post(url, json={"contents": [{"parts": [{"text": prompt}]}]}, timeout=120)
            if r.status_code == 200:
                content = r.json()['candidates'][0]['content']['parts'][0]['text']
                # Возвращаем текст с меткой начала
                return header_marker + content
            
            # Добавляем 500 и 503 в список триггеров для смены модели
            if r.status_code in [429, 404, 400, 500, 503]:
                tbox_log(f"Модель {CURRENT_MODEL} выдала {r.status_code}. Попытка {attempt+1}/3", META, "WARN", CONF)
                
                if attempt == 2: # Если это была последняя попытка - меняем модель
                    new_m = find_working_model()
                    if new_m:
                        CURRENT_MODEL = new_m
                        # Не выходим, пробуем еще раз уже с новой моделью
            
            else:
                tbox_log(f"API Error {r.status_code}", META, "ERROR", CONF)
        except Exception as e:
            tbox_log(f"Ошибка сети: {str(e)[:50]}", META, "ERROR", CONF)
        
        time.sleep(15) # Увеличиваем паузу при ошибке

    # Если всё-таки не удалось перевести
    return f"\n\n[!!! КРИТИЧЕСКАЯ ОШИБКА ЧАНКА {part_num} !!!]\n[СОРС ТЕКСТА]:\n{chunk[:200]}..."
    
def main():
    if not CONF:
        tbox_log("FATAL: config.txt не найден.", META, "ERROR")
        sys.exit(1)

    in_dir, out_dir, arh_dir = CONF.get('TEMP_TXT_DIR'), CONF.get('DOC_TRANSLATED'), CONF.get('ARH_TXT')
    
    files = glob.glob(os.path.join(in_dir, "*.txt"))
    if not files:
        tbox_log("В 02_TXT пусто.", META, "INFO", CONF)
        return

    target = max(files, key=os.path.getmtime)
    file_name = os.path.basename(target)
    author = file_name.split('-')[1].replace('_', ' ') if '-' in file_name else "Раввин"
    
    tbox_log(f"СТАРТ: {file_name} | Модель: {CURRENT_MODEL}", META, "START", CONF)
    
    with open(target, 'r', encoding='utf-8') as f:
        text = f.read()

    chunks = [text[i:i+CHUNK_SIZE] for i in range(0, len(text), CHUNK_SIZE)]
    doc = Document()
    doc.add_heading(f"Лекция: {author}", 0)

    for i, chunk in enumerate(chunks, 1):
        tbox_log(f"Часть {i}/{len(chunks)} ({CURRENT_MODEL})", META, "INFO", CONF)
        res = translate_chunk(chunk, i, len(chunks), author, '-s' in sys.argv)
        doc.add_paragraph(res)
        time.sleep(12) 

    os.makedirs(out_dir, exist_ok=True); os.makedirs(arh_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"Ready_{file_name.replace('.txt', '.docx')}")
    doc.save(out_path)
    shutil.move(target, os.path.join(arh_dir, file_name))
    tbox_log(f"ГОТОВО: {out_path}", META, "DONE", CONF)

if __name__ == "__main__":
    main()