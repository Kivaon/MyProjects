#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
TBox Configuration Loader
Загружает конфигурацию из _config/tconfig.txt
"""

import os
import re

def load_abox_config():
    """Загружает конфигурацию ABOX из _config/tconfig.txt"""
    config_path = os.path.join(os.path.dirname(__file__), '_config', 'tconfig.txt')
    
    if not os.path.exists(config_path):
        return None
    
    config = {}
    base_dir = None
    
    with open(config_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            
            if '=' in line:
                key, value = line.split('=', 1)
                key = key.strip()
                value = value.strip()
                
                # Подстановка переменных
                if '${BASE_DIR}' in value:
                    if not base_dir:
                        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                    value = value.replace('${BASE_DIR}', base_dir)
                
                config[key] = value
    
    return config

def load_local_config():
    """Алиас для load_abox_config для совместимости"""
    return load_abox_config()

def tbox_log(message, meta, level, conf=None):
    """Логирование сообщений"""
    timestamp = meta.get('date', 'unknown')
    name = meta.get('name', 'unknown')
    version = meta.get('version', 'unknown')
    
    print(f"[{timestamp}] [{name} {version}] [{level.upper()}] {message}")

def tbox_save_to_docx(content, filepath, title=None):
    """Сохранение текста в DOCX"""
    try:
        from docx import Document
        
        doc = Document()
        if title:
            doc.add_heading(title, 0)
        
        # Разбиваем текст на параграфы
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(filepath)
        return True
    except ImportError:
        print("ВНИМАНИЕ: python-docx не установлен. Установите: pip install python-docx")
        return False
    except Exception as e:
        print(f"Ошибка сохранения DOCX: {e}")
        return False

def tbox_chunk_text(text, max_chars=8000):
    """Разбивает текст на чанки"""
    if len(text) <= max_chars:
        return [text]
    
    chunks = []
    current_chunk = ""
    
    lines = text.split('\n')
    for line in lines:
        if len(current_chunk) + len(line) + 1 <= max_chars:
            current_chunk += line + '\n'
        else:
            if current_chunk:
                chunks.append(current_chunk.rstrip())
            current_chunk = line + '\n'
    
    if current_chunk:
        chunks.append(current_chunk.rstrip())
    
    return chunks
