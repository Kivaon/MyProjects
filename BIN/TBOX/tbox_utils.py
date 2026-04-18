import os, sys, re, requests, time
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- MANIFEST ---
# Previous: v2.9.bidi (2026-01-25) - Enhanced BIDI support
VERSION = "v3.0.enhanced-logging"
DATE    = "2026-04-18"
NAME    = os.path.basename(__file__)
META    = {"name": NAME, "version": VERSION}

def tbox_log(message, script_meta, level="INFO", conf=None):
    now = datetime.now()
    # ТЕПЕРЬ ТУТ И ИМЯ СКРИПТА, И ВЕРСИЯ
    version_str = script_meta.get('version', '?.?')
    tag = f"[{script_meta['name']} {version_str}]" 
    
    time_s = now.strftime('%H:%M:%S')
    time_f = now.strftime('%Y-%m-%d %H:%M:%S')
    msg = f"[{time_s}] {tag} [{level}] {message}"
    print(msg)
    
    if conf and 'LOG_FILE' in conf:
        try:
            # Разворачиваем путь к логу
            log_file = conf['LOG_FILE']
            if 'BASE_DIR' in conf and not log_file.startswith('/'):
                # Относительный путь - добавляем BASE_DIR
                if log_file.startswith('BIN/'):
                    log_file = os.path.join(conf['BASE_DIR'], log_file)
                else:
                    log_file = os.path.join(conf['BASE_DIR'], log_file)
            
            # Убеждаемся, что директория для логов существует
            log_dir = os.path.dirname(log_file)
            if log_dir and not os.path.exists(log_dir):
                os.makedirs(log_dir, exist_ok=True)
                
            with open(log_file, "a", encoding="utf-8") as f:
                f.write(f"[{time_f}] {tag} [{level}] {message}\n")
        except: pass
        
def load_local_config():
    """Загружает конфигурацию из _config/tconfig.txt"""
    config_path = os.path.join(os.path.dirname(__file__), '..', '_config', 'tconfig.txt')
    
    if not os.path.exists(config_path):
        print("КРИТИЧЕСКАЯ ОШИБКА: _config/tconfig.txt не найден.")
        return None
    
    conf = {}
    base_dir = None
    
    with open(config_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.split('#')[0].strip()
            if '=' in line:
                k, v = line.split('=', 1)
                k = k.strip()
                v = v.strip()
                conf[k] = v
                # Сохраняем BASE_DIR для подстановки
                if k == 'BASE_DIR':
                    base_dir = v
    
    # Разворачиваем переменные ${BASE_DIR}
    if base_dir:
        for k in conf:
            if isinstance(conf[k], str) and '${BASE_DIR}' in conf[k]:
                conf[k] = conf[k].replace('${BASE_DIR}', base_dir)
    
    return conf

def load_abox_config():
    """Загружает конфигурацию ABOX из _config/aconfig.txt"""
    config_path = os.path.join(os.path.dirname(__file__), '..', '_config', 'aconfig.txt')
    
    if not os.path.exists(config_path):
        print("КРИТИЧЕСКАЯ ОШИБКА: _config/aconfig.txt не найден.")
        return None
    
    conf = {}
    base_dir = None
    
    with open(config_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.split('#')[0].strip()
            if '=' in line:
                k, v = line.split('=', 1)
                k = k.strip()
                v = v.strip()
                conf[k] = v
                # Сохраняем BASE_DIR для подстановки
                if k == 'BASE_DIR':
                    base_dir = v
    
    # Разворачиваем переменные ${BASE_DIR}
    if base_dir:
        for k in conf:
            if isinstance(conf[k], str) and '${BASE_DIR}' in conf[k]:
                conf[k] = conf[k].replace('${BASE_DIR}', base_dir)
    
    return conf

def tbox_chunk_text(text, max_chars=11000):
    chunks = []
    current_chunk = []
    current_length = 0
    for line in text.split('\n'):
        if current_length + len(line) > max_chars:
            chunks.append("\n".join(current_chunk))
            current_chunk = [line]; current_length = len(line)
        else:
            current_chunk.append(line); current_length += len(line)
    if current_chunk: chunks.append("\n".join(current_chunk))
    return chunks

def tbox_save_to_docx(text, file_path, title="Lecture"):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    # Настройка полей
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2)
        s.top_margin = s.bottom_margin = Cm(1.5)

    for block in text.split('\n'):
        block = block.strip()
        if not block: continue
        
        # --- ДЕТЕКТОР ДЛЯ КАЖДОГО АБЗАЦА ---
        # Проверяем, есть ли иврит именно в этом блоке текста
        is_hebrew = bool(re.search(r'[\u0590-\u05FF]', block))
        
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        
        # 1. Настройка выравнивания (по обоим краям)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both') 
        pPr.append(jc)

        # 2. ПРИМЕНЯЕМ RTL ТОЛЬКО ЕСЛИ ЭТО ИВРИТ
        if is_hebrew:
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT # На всякий случай дублируем
        else:
            # Для русского/английского убираем BiDi флаг
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Обработка заголовка
        if block.startswith('#'):
            run = p.add_run(block.lstrip('#').strip())
            run.bold = True
            run.font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Обработка обычного текста с Markdown разметкой
            parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', block)
            for part in parts:
                if not part: continue
                clean_part = part.replace('***','').replace('**','').replace('*','')
                run = p.add_run(clean_part)
                
                # Шрифт и RTL для конкретного фрагмента
                if is_hebrew:
                    run.font.name = 'Arial'
                    rPr = run._r.get_or_add_rPr()
                    rtl_el = OxmlElement('w:rtl')
                    rtl_el.set(qn('w:val'), '1')
                    rPr.append(rtl_el)
                else:
                    run.font.name = 'Times New Roman'
                
                if '***' in part: (run.bold, run.italic) = (True, True)
                elif '**' in part: run.bold = True
                elif '*' in part: run.italic = True
                run.font.size = Pt(12)
    
    # Исправляем compatibility mode для современного Word
    try:
        from docx.oxml.ns import qn
        from lxml import etree
        
        settings = doc.settings._element
        compat = settings.find(qn('w:compat'))
        if compat is None:
            compat = etree.SubElement(settings, qn('w:compat'))
        
        for setting in compat.findall(qn('w:compatSetting')):
            if setting.get(qn('w:name')) == 'compatibilityMode':
                compat.remove(setting)
        
        new_setting = etree.SubElement(compat, qn('w:compatSetting'))
        new_setting.set(qn('w:name'), 'compatibilityMode')
        new_setting.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
        new_setting.set(qn('w:val'), '16')
    except:
        pass
                
    doc.save(file_path)