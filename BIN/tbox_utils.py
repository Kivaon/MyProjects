import os, sys, re, requests, time
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- MANIFEST ---
VERSION = "v2.9.bidi"
DATE    = "2026-01-25"
NAME    = os.path.basename(__file__)
META    = {"name": NAME, "version": VERSION}

def tbox_log(message, script_meta, level="INFO", conf=None):
    now = datetime.now()
    # ТЕПЕРЬ ТУТ И ИМЯ СКРИПТА, И ВЕРСИЯ
    version_str = script_meta.get('version', '?.?')
    tag = f"[{script_meta['name']} {version_str}]" 
    
    time_s = now.strftime('%H:%M:%S')
    time_f = now.strftime('%Y-%m-%d %H:%M:%S')
    msg = f"[{time_s}] {tag} [{level}] {message}"
    print(msg)
    
    if conf and 'LOG_FILE' in conf:
        try:
            # Убеждаемся, что директория для логов существует
            log_dir = os.path.dirname(conf['LOG_FILE'])
            if log_dir and not os.path.exists(log_dir):
                os.makedirs(log_dir, exist_ok=True)
                
            with open(conf['LOG_FILE'], "a", encoding="utf-8") as f:
                f.write(f"[{time_f}] {tag} [{level}] {message}\n")
        except: pass
        
def load_local_config():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    config_path = os.path.join(base_dir, "config.txt")
    conf = {}
    if not os.path.exists(config_path): return None
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.split('#')[0].strip()
                if '=' in line:
                    k, v = line.split('=', 1)
                    conf[k.strip()] = v.strip()
        actual_base = conf.get('BASE_DIR', base_dir)
        for k in conf:
            if isinstance(conf[k], str) and '${BASE_DIR}' in conf[k]:
                conf[k] = conf[k].replace('${BASE_DIR}', actual_base)
        return conf
    except: return None

def tbox_chunk_text(text, max_chars=11000):
    chunks = []
    current_chunk = []
    current_length = 0
    for line in text.split('\n'):
        if current_length + len(line) > max_chars:
            chunks.append("\n".join(current_chunk))
            current_chunk = [line]; current_length = len(line)
        else:
            current_chunk.append(line); current_length += len(line)
    if current_chunk: chunks.append("\n".join(current_chunk))
    return chunks

def tbox_save_to_docx(text, file_path, title="Lecture"):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    # Настройка полей
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2)
        s.top_margin = s.bottom_margin = Cm(1.5)

    for block in text.split('\n'):
        block = block.strip()
        if not block: continue
        
        # --- ДЕТЕКТОР ДЛЯ КАЖДОГО АБЗАЦА ---
        # Проверяем, есть ли иврит именно в этом блоке текста
        is_hebrew = bool(re.search(r'[\u0590-\u05FF]', block))
        
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        
        # 1. Настройка выравнивания (по обоим краям)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both') 
        pPr.append(jc)

        # 2. ПРИМЕНЯЕМ RTL ТОЛЬКО ЕСЛИ ЭТО ИВРИТ
        if is_hebrew:
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT # На всякий случай дублируем
        else:
            # Для русского/английского убираем BiDi флаг
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Обработка заголовка
        if block.startswith('#'):
            run = p.add_run(block.lstrip('#').strip())
            run.bold = True
            run.font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Обработка обычного текста с Markdown разметкой
            parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', block)
            for part in parts:
                if not part: continue
                clean_part = part.replace('***','').replace('**','').replace('*','')
                run = p.add_run(clean_part)
                
                # Шрифт и RTL для конкретного фрагмента
                if is_hebrew:
                    run.font.name = 'Arial'
                    rPr = run._r.get_or_add_rPr()
                    rtl_el = OxmlElement('w:rtl')
                    rtl_el.set(qn('w:val'), '1')
                    rPr.append(rtl_el)
                else:
                    run.font.name = 'Times New Roman'
                
                if '***' in part: (run.bold, run.italic) = (True, True)
                elif '**' in part: run.bold = True
                elif '*' in part: run.italic = True
                run.font.size = Pt(12)
                
    doc.save(file_path)