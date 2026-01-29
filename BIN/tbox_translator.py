import os, sys, time, glob, requests, shutil, re, argparse
from datetime import datetime
from docx import Document   
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tbox_utils import tbox_log

# --- MANIFEST ---
VERSION = "v6.00.hybrid_pro"
DATE    = "2026-01-27"
NAME    = os.path.basename(__file__)
META    = {"name": NAME, "version": VERSION}

def load_tbox_config():
    script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    config_path = os.path.join(script_dir, "config.txt")
    conf = {}
    if not os.path.exists(config_path): return None
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.split('#')[0].strip()
                if '=' in line:
                    key, val = line.split('=', 1)
                    conf[key.strip()] = val.strip()
        actual_base = conf.get('BASE_DIR', script_dir)
        for key in conf:
            if '${BASE_DIR}' in conf[key]:
                conf[key] = conf[key].replace('${BASE_DIR}', actual_base)
        return conf
    except Exception as e:
        print(f"Ошибка в load_tbox_config: {e}")
        return None

def load_prompts(conf):
    """Загрузка справочника промптов из файла"""
    prompts_dir = conf.get('PROMPTS_DIR', '06_PROMPTS')
    script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    prompts_path = os.path.join(script_dir, prompts_dir, 'prompts.md')
    prompts = {
        'TORAH': "Ты — редактор лекций Рава {author}. Переведи часть {part_num}/{total_parts} на русский.\nСТРОГО СОХРАНЯЙ РАЗМЕТКУ: '#' для заголовков и '**' для выделений.\nСтиль: возвышенный, точный.\nТЕКСТ:\n{chunk}",
        'FICTION': "Переведи художественный текст часть {part_num}/{total_parts} на русский.\nСОХРАНЯЙ стилистику и тон оригинала.\nРазметка: '#' для заголовков, '**' для выделений.\nТЕКСТ:\n{chunk}",
        'GENERIC': "Переведи текст часть {part_num}/{total_parts} на русский.\nТЕКСТ:\n{chunk}"
    }
    if os.path.exists(prompts_path):
        try:
            with open(prompts_path, 'r', encoding='utf-8') as f:
                content = f.read()
            # Парсинг по заголовкам # CODE
            sections = re.split(r'^# (\w+)$', content, flags=re.MULTILINE)
            for i in range(1, len(sections), 2):
                code = sections[i].strip()
                text = sections[i+1].strip()
                prompts[code] = text
        except Exception as e:
            tbox_log(f"Ошибка загрузки промптов: {e}", META, "ERROR", conf)
    else:
        tbox_log(f"Файл промптов не найден, используем встроенные: {prompts_path}", META, "WARN", conf)
    return prompts

def get_api_ver(model_name):
    """Определяет API версию по названию модели"""
    v_match = re.search(r'(\d+)', model_name)
    v_major = int(v_match.group(1)) if v_match else 1
    if v_major >= 2 or any(x in model_name.lower() for x in ['exp', 'beta']):
        return "v1beta"
    return "v1"

def find_working_model(api_key, conf):
    """Автоматический подбор рабочей модели при сбое основной"""
    tbox_log("ПОИСК АЛЬТЕРНАТИВНОЙ МОДЕЛИ...", META, "WARN", conf)
    
    def v_score(name):
        v = re.search(r'(\d+\.?\d*)', name)
        num = float(v.group(1)) if v else 0.0
        score = num + (0.5 if "flash" in name.lower() else 0)
        if any(x in name.lower() for x in ["vision", "pro", "exp"]): score -= 1
        return score

    for api_v in ["v1beta", "v1"]:
        try:
            url = f"https://generativelanguage.googleapis.com/{api_v}/models?key={api_key}"
            r = requests.get(url, timeout=10)
            if r.status_code == 200:
                models = [m['name'].split('/')[-1] for m in r.json().get('models', []) 
                          if 'generateContent' in m.get('supportedGenerationMethods', [])]
                models.sort(key=v_score, reverse=True)
                
                for model in models:
                    test_url = f"https://generativelanguage.googleapis.com/{api_v}/models/{model}:generateContent?key={api_key}"
                    try:
                        tr = requests.post(test_url, json={"contents": [{"parts": [{"text": "test"}]}]}, timeout=5)
                        if tr.status_code == 200:
                            tbox_log(f"Найдена модель: {model}", META, "INFO", conf)
                            return model
                    except:
                        continue
        except:
            continue
    return None

def translate_md_chunk(chunk, part_num, total_parts, author, conf, prompts, prompt_code, include_original):
    """Перевод MD-текста с retry-логикой и переключением модели"""
    global CURRENT_MODEL
    api_key = conf.get('API_KEY', '').split('#')[0].strip()
    
    # Получить промпт из справочника
    base_prompt = prompts.get(prompt_code, prompts.get('GENERIC', 'Переведи на русский: {chunk}'))
    
    # Подставить плейсхолдеры
    prompt = base_prompt.format(
        author=author,
        part_num=part_num,
        total_parts=total_parts,
        chunk=chunk
    )
    
    # Если include_original, добавить инструкцию о цитатах
    if include_original:
        prompt += "\n\nВКЛЮЧАЙ ОРИГИНАЛЬНЫЙ ТЕКСТ ЦИТАТ В СКОБКАХ ПОСЛЕ ПЕРЕВОДА."

    while True:
        # Маркер начала чанка с текущей моделью
        header_marker = f"\n\n--- [PART {part_num}/{total_parts} | MODEL: {CURRENT_MODEL}] ---\n"
        
        # 3 попытки для текущей модели
        for attempt in range(3):
            try:
                api_ver = get_api_ver(CURRENT_MODEL)
                url = f"https://generativelanguage.googleapis.com/{api_ver}/models/{CURRENT_MODEL}:generateContent?key={api_key}"
                
                tbox_log(f"Попытка {attempt+1}/3: {CURRENT_MODEL}", META, "INFO", conf)
                
                r = requests.post(url, json={"contents": [{"parts": [{"text": prompt}]}]}, timeout=120)
                
                if r.status_code == 200:
                    content = r.json()['candidates'][0]['content']['parts'][0]['text']
                    tbox_log(f"Успех: часть {part_num}/{total_parts}", META, "INFO", conf)
                    return header_marker + content
                
                # Дифференцированная обработка ошибок
                if r.status_code == 404:
                    # Модель не существует/удалена — менять сразу
                    tbox_log(f"Модель {CURRENT_MODEL} не найдена (404)", META, "WARN", conf)
                    new_model = find_working_model(api_key, conf)
                    if new_model and new_model != CURRENT_MODEL:
                        CURRENT_MODEL = new_model
                        tbox_log(f"Переключились на: {CURRENT_MODEL}", META, "INFO", conf)
                        break
                    else:
                        return f"\n\n[ОШИБКА 404: Модель не найдена]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                
                elif r.status_code == 400:
                    # Bad Request — ошибка в запросе/промпте, не в модели
                    tbox_log(f"Ошибка 400 Bad Request: проблема с промптом", META, "ERROR", conf)
                    return f"\n\n[ОШИБКА 400: Bad Request]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                
                elif r.status_code in [429, 500, 503]:
                    # Rate Limit, Server Error, Service Unavailable — повторять
                    tbox_log(f"Ошибка {r.status_code} на попытке {attempt+1}/3", META, "WARN", conf)
                    
                    if attempt == 2:  # На последней попытке ищем замену
                        new_model = find_working_model(api_key, conf)
                        if new_model and new_model != CURRENT_MODEL:
                            CURRENT_MODEL = new_model
                            tbox_log(f"Переключились на: {CURRENT_MODEL}", META, "INFO", conf)
                            break
                        else:
                            return f"\n\n[ОШИБКА {r.status_code}: Все модели недоступны]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                
                else:
                    # Неизвестная ошибка — обрабатывать как временная
                    tbox_log(f"API Error {r.status_code} на попытке {attempt+1}/3", META, "WARN", conf)
                    
                    if attempt == 2:  # На последней попытке ищем замену
                        new_model = find_working_model(api_key, conf)
                        if new_model and new_model != CURRENT_MODEL:
                            CURRENT_MODEL = new_model
                            tbox_log(f"Переключились на: {CURRENT_MODEL}", META, "INFO", conf)
                            break
                        else:
                            return f"\n\n[ОШИБКА {r.status_code}: Все модели недоступны]\n[ИСХОДНИК]:\n{chunk[:300]}..."
                    
            except Exception as e:
                tbox_log(f"Сетевая ошибка: {str(e)[:50]}", META, "ERROR", conf)
            
            if attempt < 2:
                time.sleep(15)  # Пауза перед повтором
        else:
            # Цикл завершился без break — все 3 попытки провалились и не было смены модели
            return f"\n\n[КРИТИЧЕСКАЯ ОШИБКА ЧАНКА {part_num}]\n[ИСХОДНИК]:\n{chunk[:300]}..."

def render_md_to_docx(md_text, doc):
    """Отрисовка Markdown в параграфы Word"""
    for line in md_text.split('\n'):
        line = line.strip()
        if not line: continue
        p = doc.add_paragraph()
        if line.startswith('#'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.lstrip('#').strip())
            run.bold = True; run.font.size = Pt(14)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                run = p.add_run(part.replace('**', ''))
                run.font.name = 'Times New Roman'; run.font.size = Pt(12)
                if part.startswith('**'): run.bold = True

def main():
    global CURRENT_MODEL
    
    # Парсинг аргументов
    parser = argparse.ArgumentParser(description='TBox Translator')
    parser.add_argument('file', nargs='?', help='Путь к файлу или имя файла')
    parser.add_argument('-txt', action='store_true', help='Использовать последний txt из TXT_RAW')
    parser.add_argument('-md', action='store_true', help='Использовать последний md из MD_DIR')
    parser.add_argument('-s', action='store_true', help='Включать оригинал цитат')
    parser.add_argument('-p', '--prompt', type=str, help='Код промпта (TORAH, FICTION, GENERIC)', dest='prompt')
    args = parser.parse_args()
    
    CONF = load_tbox_config()
    if not CONF: 
        print("Ошибка: Конфиг не найден")
        return
    
    # Загрузка промптов
    PROMPTS = load_prompts(CONF)
    
    # Выбор промпта
    prompt_code = args.prompt or CONF.get('DEFAULT_PROMPT', 'TORAH')
    if prompt_code not in PROMPTS:
        prompt_code = 'GENERIC'
    tbox_log(f"Выбран промпт: {prompt_code}", META, "INFO", CONF)
    
    # Инициализация глобальной модели
    CURRENT_MODEL = CONF.get('MODEL_GEMINI', 'gemini-2.0-flash').strip()
    
    # Пути из конфига
    txt_raw_dir = CONF.get('TXT_RAW')           # 02_TXT/raw
    md_dir = CONF.get('MD_DIR', '02_TXT/MD')    # Для md файлов
    out_dir = CONF.get('DOC_TRANSLATED')        # 03_DOC/TRANSLATED
    arh_dir = CONF.get('ARH_TXT')               # 05_ARH/TXT

    # --- ВЫБОР ФАЙЛА ---
    target = None
    if args.file:
        if os.path.exists(args.file):
            target = args.file
        elif args.file.endswith('.txt'):
            potential_path = os.path.join(txt_raw_dir, args.file)
            if os.path.exists(potential_path):
                target = potential_path
        elif args.file.endswith('.md'):
            potential_path = os.path.join(md_dir, args.file)
            if os.path.exists(potential_path):
                target = potential_path
    
    if not target:
        if args.md:
            # Последний md из MD_DIR
            files = glob.glob(os.path.join(md_dir, "*.md"))
            if files:
                target = max(files, key=os.path.getmtime)
        elif args.txt or not (args.md or args.txt):
            # Последний txt из TXT_RAW (по умолчанию)
            files = glob.glob(os.path.join(txt_raw_dir, "*.txt"))
            if files:
                target = max(files, key=os.path.getmtime)
    
    if not target:
        tbox_log("Нет подходящих файлов для перевода", META, "INFO", CONF)
        return

    file_name = os.path.basename(target)
    author = file_name.split('-')[1].replace('_', ' ') if '-' in file_name else "Раввин"
    
    tbox_log(f"СТАРТ ПЕРЕВОДА: {file_name}", META, "START", CONF)

    with open(target, 'r', encoding='utf-8') as f:
        full_text = f.read()

    # Делим на чанки
    chunks = [full_text[i:i+10000] for i in range(0, len(full_text), 10000)]
    
    doc = Document()
    # Настройка страницы
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    doc.add_heading(f"Перевод лекции: {author}", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, chunk in enumerate(chunks, 1):
        tbox_log(f"Перевод чанка {i}/{len(chunks)}...", META, "INFO", CONF)
        translated_md = translate_md_chunk(chunk, i, len(chunks), author, CONF, PROMPTS, prompt_code, args.s)
        render_md_to_docx(translated_md, doc)
        if i < len(chunks):
            time.sleep(12)

    # Сохранение
    os.makedirs(out_dir, exist_ok=True)
    res_name = f"Ready_{file_name.replace('.txt', '.docx').replace('.md', '.docx')}"
    res_path = os.path.join(out_dir, res_name)
    doc.save(res_path)
    
    # Архивируем
    os.makedirs(arh_dir, exist_ok=True)
    try:
        shutil.move(target, os.path.join(arh_dir, f"TR_DONE_{file_name}"))
        tbox_log(f"ГОТОВО: {res_name}", META, "DONE", CONF)
    except Exception as e:
        tbox_log(f"Файл сохранен, архивирование ошибка: {e}", META, "WARNING", CONF)

if __name__ == "__main__":
    main()